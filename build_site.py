#!/usr/bin/env python3
"""Génère le site statique pour la bibliothèque de talks Satipanya."""

import hashlib
import json
import re
import shutil
from pathlib import Path
from html import escape as h
from fpdf import FPDF

# ── Configuration ──────────────────────────────────────────────

PROJECT_DIR = Path(__file__).parent
CATALOG_PATH = PROJECT_DIR / "catalog.json"
METADATA_DIR = PROJECT_DIR / "metadata"
ARTICLES_DIR = PROJECT_DIR / "articles"
COVERS_DIR  = PROJECT_DIR / "covers"
SITE_DIR    = PROJECT_DIR / "site"

SITE_TITLE = "Satipanya Dharma Library"
SITE_TAGLINE = "Dharma talks, guided meditations, essays and teachings from Satipanya Buddhist Retreat"
FEEDS_BASE_URL = "https://www.enpleineconscience.ch/satipanya"

# Préfixe URL pour héberger le site dans un sous-dossier (ex: "/satipanya")
# Laisser vide "" pour un hébergement à la racine
SITE_BASE_PATH = "/satipanya"

# Ordre d'affichage des feeds sur la page d'accueil
AUDIO_FEED_ORDER = [
    "dharma-talks",
    "youtube-talks",
    "noirins-teachings",
    "dhammabytes",
    "guided-meditations",
    "foundation-course",
    "international-talks",
]

TEXT_FEED_ORDER = [
    "bhante-essays",
    "noirin-essays",
    "tips-of-the-day",
    "retreat-talks",
]

FEED_ORDER = AUDIO_FEED_ORDER + TEXT_FEED_ORDER

# Feeds affichés du plus récent au plus ancien (comme un blog/fil d'actualité)
FEEDS_NEWEST_FIRST = {"dharma-talks", "youtube-talks", "noirins-teachings"}

FEED_EMOJI = {
    "guided-meditations": "🧘",
    "foundation-course": "📚",
    "dhammabytes": "💎",
    "dharma-talks": "🪷",
    "noirins-teachings": "🌿",
    "international-talks": "🌍",
    "youtube-talks": "🎬",
    "bhante-essays": "✍️",
    "noirin-essays": "📝",
    "tips-of-the-day": "💡",
    "retreat-talks": "🏔️",
}

# ── Utilitaires ────────────────────────────────────────────────

def format_duration(seconds):
    """Formate des secondes en HH:MM:SS ou MM:SS."""
    if not seconds:
        return ""
    s = int(seconds)
    h_val, s = divmod(s, 3600)
    m, s = divmod(s, 60)
    if h_val:
        return f"{h_val}:{m:02d}:{s:02d}"
    return f"{m}:{s:02d}"


def ep_stem(ep):
    """Extrait le nom de fichier (sans extension) du transcript_path ou du champ stem."""
    if ep.get("stem"):
        return ep["stem"]
    tp = ep.get("transcript_path")
    if tp:
        return Path(tp).stem
    return None


def is_text_episode(ep):
    """True si l'épisode est du contenu écrit (pas audio)."""
    return ep.get("content_type") == "text"


def is_text_feed(fdata):
    """True si la collection est du contenu écrit."""
    return fdata.get("content_type") == "text"


def format_reading_time(minutes):
    """Formate un temps de lecture en 'N min read'."""
    if not minutes:
        return ""
    return f"{int(minutes)} min read"


def base(path):
    """Préfixe un chemin absolu avec SITE_BASE_PATH (ex: /satipanya/style.css)."""
    return f"{SITE_BASE_PATH}{path}"


# Mois anglais → numéro (pour parser les dates dans les titres d'épisodes)
_MONTHS = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
    "january": 1, "february": 2, "march": 3, "april": 4, "june": 6,
    "july": 7, "august": 8, "september": 9, "october": 10, "november": 11, "december": 12,
}


def _season_sort_key(season):
    """Clé de tri pour ordonner les saisons par date la plus récente.

    Analyse les noms de saisons et titres d'épisodes pour extraire des dates.
    Retourne (max_year, max_month, max_day) — les saisons sans date donnent (0,0,0).

    Si le nom de la saison contient une année explicite (ex: "Gaia House 2011"),
    celle-ci est utilisée directement — les titres d'épisodes ne sont pas consultés
    (ils contiennent parfois des chiffres parasites comme des tailles de fichier).
    """
    name = season.get("name", "")

    # "Most Recent" = contenu le plus récent, toujours en tête
    if "most recent" in name.lower():
        return (9999, 12, 31)

    # Si le nom de la saison contient une année, l'utiliser directement
    name_years = [int(y) for y in re.findall(r'20[0-2]\d', name)]
    if name_years:
        return (max(name_years), 0, 0)

    # Sinon, chercher des dates dans les titres d'épisodes
    best = (0, 0, 0)
    for ep in season.get("episodes", []):
        text = ep.get("title", "")
        # Pattern: "2024 April.24", "2025 Jan 12", "2022 Dec 11"
        for m in re.finditer(r'(20[0-2]\d)\s+(\w+)\.?\s+(\d{1,2})', text):
            y = int(m.group(1))
            mon = _MONTHS.get(m.group(2).lower().rstrip('.'), 0)
            if mon == 0:
                continue  # Le mot après l'année n'est pas un mois → ignorer
            d = int(m.group(3))
            best = max(best, (y, mon, d))
        # Pattern: "2024 April" ou "2025 Jan" (sans jour)
        for m in re.finditer(r'(20[0-2]\d)\s+(\w+)', text):
            y = int(m.group(1))
            mon = _MONTHS.get(m.group(2).lower().rstrip('.'), 0)
            if mon:
                best = max(best, (y, mon, 0))
    return best


def ep_url(feed_slug, stem):
    """URL relative vers la page d'un épisode."""
    return base(f"/{feed_slug}/{stem}.html")


def load_metadata(feed_slug, stem):
    """Charge le fichier metadata JSON d'un épisode."""
    path = METADATA_DIR / feed_slug / f"{stem}.json"
    if path.exists():
        with open(path) as f:
            return json.load(f)
    return {}


def load_article(feed_slug, stem):
    """Charge le texte beautifié d'un épisode."""
    path = ARTICLES_DIR / feed_slug / f"{stem}.txt"
    if path.exists():
        return path.read_text(encoding="utf-8").strip()
    return None


def article_to_html(text):
    """Convertit un article texte en HTML avec paragraphes et italiques."""
    paragraphs = text.split("\n\n")
    html_parts = []
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        # Convertir *italiques* en <em>
        p = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', h(p))
        # Restaurer les sauts de ligne simples à l'intérieur d'un paragraphe
        p = p.replace("\n", "<br>")
        html_parts.append(f"<p>{p}</p>")
    return "\n".join(html_parts)


def keyword_tags_html(keywords):
    """Génère les tags de mots-clés."""
    if not keywords:
        return ""
    tags = "".join(
        f'<a href="{base("/search.html")}?q={h(kw)}" class="tag">{h(kw)}</a>'
        for kw in keywords
    )
    return f'<div class="tags">{tags}</div>'


# ── PDF ────────────────────────────────────────────────────────

FONT_DIR = Path("/usr/share/fonts/truetype/noto")

class TranscriptPDF(FPDF):
    """PDF avec en-tête et pied de page Satipanya."""

    def __init__(self, title, speaker, feed_name):
        super().__init__()
        self._title = title
        self._speaker = speaker
        self._feed_name = feed_name
        self.add_font("NotoSerif", "", str(FONT_DIR / "NotoSerif-Regular.ttf"))
        self.add_font("NotoSerif", "I", str(FONT_DIR / "NotoSerif-Italic.ttf"))
        self.add_font("NotoSerif", "B", str(FONT_DIR / "NotoSerif-Bold.ttf"))
        self.set_auto_page_break(auto=True, margin=25)

    def header(self):
        if self.page_no() == 1:
            return  # première page a son propre en-tête
        self.set_font("NotoSerif", "I", 8)
        self.set_text_color(120, 113, 108)
        self.cell(0, 8, self._title, align="L")
        self.ln(12)

    def footer(self):
        self.set_y(-20)
        self.set_font("NotoSerif", "I", 8)
        self.set_text_color(120, 113, 108)
        self.cell(0, 8, f"Satipanya Buddhist Retreat — {self._feed_name}", align="L")
        self.cell(0, 8, str(self.page_no()), align="R")


def generate_pdf(output_path, title, speaker, feed_name, duration_str, article_text):
    """Génère un PDF élégant pour un transcript beautifié."""
    pdf = TranscriptPDF(title, speaker, feed_name)
    pdf.add_page()

    # ── Page de titre intégrée ──
    pdf.ln(15)
    pdf.set_font("NotoSerif", "B", 22)
    pdf.set_text_color(45, 41, 38)
    pdf.multi_cell(0, 11, title, align="L")
    pdf.ln(4)
    pdf.set_font("NotoSerif", "", 11)
    pdf.set_text_color(120, 113, 108)
    pdf.cell(0, 7, f"{speaker}  ·  {feed_name}  ·  {duration_str}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    # Ligne décorative
    pdf.set_draw_color(184, 134, 11)  # saffron
    pdf.set_line_width(0.5)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(10)

    # ── Corps du texte ──
    pdf.set_text_color(45, 41, 38)
    paragraphs = article_text.split("\n\n")
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        # Découper en segments normal / *italique*
        parts = re.split(r'(\*[^*]+\*)', para)
        for part in parts:
            if part.startswith("*") and part.endswith("*"):
                pdf.set_font("NotoSerif", "I", 10.5)
                pdf.write(6.5, part[1:-1])
            else:
                pdf.set_font("NotoSerif", "", 10.5)
                pdf.write(6.5, part)
        pdf.ln(10)

    # ── Note de fin ──
    pdf.ln(5)
    pdf.set_draw_color(184, 134, 11)
    pdf.set_line_width(0.3)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + 40, pdf.get_y())
    pdf.ln(4)
    pdf.set_font("NotoSerif", "I", 8)
    pdf.set_text_color(120, 113, 108)
    pdf.multi_cell(0, 5,
        "Transcriptions produced locally using Swiss low-carbon electricity. "
        "Corrections and rewriting by cloud-hosted AI.")

    pdf.output(str(output_path))


def generate_docx(output_path, title, speaker, feed_name, duration_str, article_text):
    """Génère un DOCX pour un transcript individuel."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.8)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ── Titre ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True

    p = doc.add_paragraph()
    run = p.add_run(f"{speaker}  ·  {feed_name}  ·  {duration_str}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)
    p.paragraph_format.space_after = Pt(10)

    # ── Ligne décorative ──
    p = doc.add_paragraph()
    run = p.add_run("─" * 30)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
    p.paragraph_format.space_after = Pt(12)

    # ── Corps du texte ──
    paragraphs = article_text.split("\n\n")
    for p_text in paragraphs:
        p_text = p_text.strip()
        if not p_text:
            continue
        p = doc.add_paragraph()
        parts = re.split(r'(\*[^*]+\*)', p_text)
        for part in parts:
            if part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                p.add_run(part.replace("\n", " "))

    # ── Note de fin ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Transcriptions produced locally using Swiss low-carbon electricity.\n"
        "Corrections and rewriting by cloud-hosted AI."
    )
    run.font.size = Pt(7)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(str(output_path))


EPUB_CSS = """\
body {
    font-family: Georgia, "Times New Roman", serif;
    font-size: 1em;
    line-height: 1.7;
    color: #1a1a1a;
    margin: 0;
    padding: 0;
}
h1 {
    font-size: 1.6em;
    font-weight: 700;
    line-height: 1.25;
    margin: 0 0 0.3em;
    color: #1a1a1a;
}
.chapter-meta {
    font-size: 0.8em;
    color: #999;
    margin-bottom: 1em;
}
.chapter-lead {
    font-style: italic;
    color: #555;
    border-left: 3px solid #e0d8cf;
    padding-left: 1em;
    margin-bottom: 1.5em;
    line-height: 1.7;
}
.chapter-lead p { margin-bottom: 0.5em; }
.chapter-body p {
    margin-bottom: 0.8em;
    text-align: justify;
}
.chapter-body em { font-style: italic; }
.colophon {
    text-align: center;
    margin-top: 3em;
    color: #999;
    font-size: 0.85em;
    line-height: 1.8;
}
"""


def generate_epub(output_path, title, speaker, feed_name, duration_str, article_text,
                  description=""):
    """Génère un EPUB pour un transcript individuel."""
    from ebooklib import epub
    book = epub.EpubBook()
    slug = Path(output_path).stem
    book.set_identifier(f"satipanya-{slug}")
    book.set_title(title)
    book.set_language("en")
    if speaker:
        book.add_author(speaker)
    book.add_metadata("DC", "publisher", "Satipanya Buddhist Retreat")

    style = epub.EpubItem(uid="style", file_name="style/default.css",
                          media_type="text/css", content=EPUB_CSS.encode())
    book.add_item(style)

    # Description en accroche
    lead_html = ""
    if description:
        desc_paras = "".join(
            f"<p>{h(p.strip())}</p>" for p in description.split("\n\n") if p.strip()
        )
        lead_html = f'<div class="chapter-lead">{desc_paras}</div>'

    body_html = article_to_html(article_text)

    ch = epub.EpubHtml(title=title, file_name="content.xhtml", lang="en")
    ch.content = f"""<html><head></head><body>
<h1>{h(title)}</h1>
<div class="chapter-meta">{h(speaker)} · {h(feed_name)} · {duration_str}</div>
{lead_html}
<div class="chapter-body">{body_html}</div>
<div class="colophon">
<p>Transcriptions produced locally using Swiss low-carbon electricity.
Corrections and rewriting by cloud-hosted AI.</p>
<p><a href="https://www.satipanya.org.uk">satipanya.org.uk</a></p>
</div>
</body></html>"""
    ch.add_item(style)
    book.add_item(ch)

    book.toc = [ch]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", ch]

    epub.write_epub(str(output_path), book, {})


# ── CSS ────────────────────────────────────────────────────────

CSS = """\
/* ── Reset ───────────────────────────────────────────── */
*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

/* ── Variables ───────────────────────────────────────── */
:root {
  --bg: #FDFBF7;
  --bg-card: #FFFFFF;
  --bg-alt: #F7F4EF;
  --bg-hero: linear-gradient(135deg, #2D2926 0%, #44403C 100%);
  --text: #2D2926;
  --text-secondary: #78716C;
  --text-inverse: #FDFBF7;
  --accent: #B8860B;
  --accent-hover: #96700A;
  --accent-light: #FEF3C7;
  --accent-subtle: #F5E6CF;
  --border: #E7E5E4;
  --border-light: #F0EEEB;
  --link: #5B7B6F;
  --link-hover: #3D5A4F;
  --shadow-sm: 0 1px 2px rgba(0,0,0,0.04);
  --shadow: 0 1px 3px rgba(0,0,0,0.06), 0 1px 2px rgba(0,0,0,0.04);
  --shadow-md: 0 4px 6px rgba(0,0,0,0.05), 0 2px 4px rgba(0,0,0,0.04);
  --radius: 10px;
  --radius-sm: 6px;
  --font-serif: 'Crimson Pro', Georgia, 'Times New Roman', serif;
  --font-sans: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
  --max-w: 1120px;
  --content-w: 700px;
}

/* ── Base ────────────────────────────────────────────── */
html { scroll-behavior: smooth; }
body {
  font-family: var(--font-sans);
  color: var(--text);
  background: var(--bg);
  line-height: 1.6;
  font-size: 15px;
  -webkit-font-smoothing: antialiased;
}
a { color: var(--link); text-decoration: none; transition: color .15s; }
a:hover { color: var(--link-hover); }
img { max-width: 100%; height: auto; display: block; }

/* ── Header ──────────────────────────────────────────── */
.site-header {
  background: var(--bg-card);
  border-bottom: 1px solid var(--border);
  position: sticky; top: 0; z-index: 100;
  backdrop-filter: blur(12px);
  background: rgba(253,251,247,0.92);
}
.header-inner {
  max-width: var(--max-w);
  margin: 0 auto;
  padding: 0 1.5rem;
  height: 56px;
  display: flex; align-items: center; justify-content: space-between;
}
.site-logo {
  font-family: var(--font-serif);
  font-size: 1.25rem;
  font-weight: 600;
  color: var(--text);
  letter-spacing: -0.01em;
}
.site-logo:hover { color: var(--accent); }
.header-nav { display: flex; align-items: center; gap: 1.5rem; }
.header-nav a { font-size: 0.875rem; color: var(--text-secondary); font-weight: 500; }
.header-nav a:hover { color: var(--accent); }
.header-search {
  display: flex; align-items: center;
  background: var(--bg-alt);
  border: 1px solid var(--border);
  border-radius: 100px;
  padding: 0.375rem 0.75rem;
  gap: 0.5rem;
  transition: border-color .15s, box-shadow .15s;
}
.header-search:focus-within {
  border-color: var(--accent);
  box-shadow: 0 0 0 3px var(--accent-light);
}
.header-search input {
  border: none; background: none; outline: none;
  font-size: 0.875rem; font-family: var(--font-sans);
  color: var(--text); width: 180px;
}
.header-search input::placeholder { color: var(--text-secondary); }
.header-search svg { flex-shrink: 0; color: var(--text-secondary); }

/* ── Hero ────────────────────────────────────────────── */
.hero {
  background: var(--bg-hero);
  color: var(--text-inverse);
  padding: 4rem 1.5rem;
  text-align: center;
}
.hero h1 {
  font-family: var(--font-serif);
  font-size: clamp(2rem, 5vw, 3rem);
  font-weight: 600;
  margin-bottom: 0.75rem;
  letter-spacing: -0.02em;
}
.hero p {
  font-size: 1.1rem;
  opacity: 0.85;
  max-width: 560px;
  margin: 0 auto 2rem;
  line-height: 1.7;
}
.hero-stats {
  display: flex; justify-content: center; gap: 2.5rem;
  font-size: 0.875rem; opacity: 0.7;
}
.hero-stats strong { display: block; font-size: 1.5rem; opacity: 1; font-weight: 700; }

/* ── Container ───────────────────────────────────────── */
.container { max-width: var(--max-w); margin: 0 auto; padding: 0 1.5rem; }

/* ── Feed Cards (homepage) ───────────────────────────── */
.feeds-section { padding: 3rem 0 4rem; }
.feeds-section h2 {
  font-family: var(--font-serif);
  font-size: 1.5rem;
  font-weight: 600;
  margin-bottom: 1.5rem;
  text-align: center;
  color: var(--text);
}
.feeds-grid {
  display: grid;
  grid-template-columns: repeat(auto-fill, minmax(320px, 1fr));
  gap: 1.25rem;
}
.feed-card {
  background: var(--bg-card);
  border: 1px solid var(--border);
  border-radius: var(--radius);
  overflow: hidden;
  transition: box-shadow .2s, transform .2s;
  display: flex; flex-direction: column;
}
.feed-card:hover {
  box-shadow: var(--shadow-md);
  transform: translateY(-2px);
}
a.feed-card { color: inherit; text-decoration: none; }
.feed-card-header {
  padding: 1.5rem 1.5rem 1rem;
  display: flex; align-items: flex-start; gap: 1rem;
}
.feed-card-emoji { font-size: 2rem; line-height: 1; flex-shrink: 0; }
.feed-card-header h3 {
  font-family: var(--font-serif);
  font-size: 1.25rem;
  font-weight: 600;
  line-height: 1.3;
}
.feed-card-header h3 a { color: var(--text); }
.feed-card-header h3 a:hover { color: var(--accent); }
.feed-card-body {
  padding: 0 1.5rem 1.5rem;
  flex: 1;
}
.feed-card-body p {
  font-size: 0.9rem;
  color: var(--text-secondary);
  line-height: 1.6;
  margin-bottom: 1rem;
}
.feed-card-meta {
  font-size: 0.8rem;
  color: var(--text-secondary);
  display: flex; flex-wrap: wrap; gap: 0.5rem 1rem;
  align-items: center;
}
.feed-card-meta span {
  display: flex; align-items: center; gap: 0.3rem;
}
.feed-card-meta .subscribe {
  color: var(--accent);
  transition: color .15s;
}
.feed-card-meta .subscribe:first-of-type {
  margin-left: auto;
}
.feed-card-meta .subscribe:hover { color: var(--accent-hover); }

/* ── Feed Page ───────────────────────────────────────── */
.feed-header {
  padding: 2.5rem 0 2rem;
  border-bottom: 1px solid var(--border);
  margin-bottom: 2rem;
}
.feed-header h1 {
  font-family: var(--font-serif);
  font-size: clamp(1.75rem, 4vw, 2.25rem);
  font-weight: 600;
  margin-bottom: 0.5rem;
}
.feed-header p {
  color: var(--text-secondary);
  max-width: var(--content-w);
  line-height: 1.7;
}
.season-section { margin-bottom: 2.5rem; }
.season-title {
  font-family: var(--font-serif);
  font-size: 1.2rem;
  font-weight: 600;
  color: var(--accent);
  margin-bottom: 1rem;
  padding-bottom: 0.5rem;
  border-bottom: 2px solid var(--accent-light);
}
.episode-list { list-style: none; }
.episode-item {
  border: 1px solid var(--border-light);
  border-radius: var(--radius-sm);
  padding: 1rem 1.25rem;
  margin-bottom: 0.5rem;
  background: var(--bg-card);
  transition: border-color .15s, box-shadow .15s;
  display: grid;
  grid-template-columns: 1fr auto;
  gap: 0.5rem 1rem;
  align-items: start;
}
.episode-item:hover {
  border-color: var(--accent-light);
  box-shadow: var(--shadow-sm);
}
.episode-title {
  font-family: var(--font-serif);
  font-size: 1.05rem;
  font-weight: 600;
  line-height: 1.4;
}
.episode-title a { color: var(--text); }
.episode-title a:hover { color: var(--accent); }
.episode-desc {
  grid-column: 1 / -1;
  font-size: 0.875rem;
  color: var(--text-secondary);
  line-height: 1.6;
}
.episode-duration {
  font-size: 0.8rem;
  color: var(--text-secondary);
  white-space: nowrap;
  padding-top: 0.15rem;
}

/* ── Episode Page ────────────────────────────────────── */
.episode-header { padding: 2.5rem 0 1.5rem; }
.episode-header h1 {
  font-family: var(--font-serif);
  font-size: clamp(1.5rem, 4vw, 2rem);
  font-weight: 600;
  line-height: 1.3;
  margin-bottom: 0.75rem;
}
.episode-meta {
  display: flex; flex-wrap: wrap; gap: 0.5rem 1.5rem;
  font-size: 0.875rem;
  color: var(--text-secondary);
  margin-bottom: 1.5rem;
}
.episode-meta span { display: flex; align-items: center; gap: 0.35rem; }

/* Audio player */
.audio-player {
  background: var(--bg-alt);
  border: 1px solid var(--border);
  border-radius: var(--radius);
  padding: 1.25rem;
  margin-bottom: 2rem;
}
.audio-player audio {
  width: 100%;
  height: 40px;
}
.youtube-embed {
  position: relative;
  padding-bottom: 56.25%;
  height: 0;
  overflow: hidden;
  border-radius: 8px;
}
.youtube-embed iframe {
  position: absolute;
  top: 0; left: 0;
  width: 100%; height: 100%;
  border-radius: 8px;
}
.audio-source {
  margin-top: 0.5rem;
  font-size: 0.75rem;
  color: var(--text-secondary);
}
.audio-source a { color: var(--text-secondary); text-decoration: underline; }

/* Description */
.episode-description {
  margin-bottom: 2rem;
  line-height: 1.8;
  color: var(--text);
}
.episode-description p { margin-bottom: 1rem; }
.episode-description p:last-child { margin-bottom: 0; }

/* Tags */
.tags { display: flex; flex-wrap: wrap; gap: 0.4rem; margin-bottom: 2rem; }
.tag {
  display: inline-block;
  background: var(--accent-light);
  color: var(--accent);
  font-size: 0.75rem;
  font-weight: 500;
  padding: 0.2rem 0.6rem;
  border-radius: 100px;
  transition: background .15s, color .15s;
}
.tag:hover { background: var(--accent); color: white; }

/* Transcript */
.transcript-section {
  border-top: 1px solid var(--border);
  padding-top: 2rem;
  margin-bottom: 3rem;
}
.transcript-section h2 {
  font-family: var(--font-serif);
  font-size: 1.25rem;
  font-weight: 600;
  margin-bottom: 1.5rem;
  color: var(--text);
}
.transcript-text {
  font-family: var(--font-serif);
  font-size: 1.1rem;
  line-height: 2;
  color: #3D3835;
  max-width: var(--content-w);
}
.transcript-text p {
  margin-bottom: 1.25rem;
  text-indent: 0;
}
.transcript-text em { font-style: italic; }
.transcript-unavailable {
  color: var(--text-secondary);
  font-style: italic;
  padding: 2rem;
  text-align: center;
  background: var(--bg-alt);
  border-radius: var(--radius);
}
.transcript-header {
  display: flex; align-items: center; justify-content: space-between;
  flex-wrap: wrap; gap: 0.75rem;
  margin-bottom: 1.5rem;
}
.transcript-header h2 { margin-bottom: 0; }
.btn-pdf {
  display: inline-flex; align-items: center; gap: 0.4rem;
  font-size: 0.8rem; font-weight: 500;
  color: var(--accent);
  background: var(--accent-light);
  padding: 0.35rem 0.85rem;
  border-radius: 100px;
  transition: background .15s, color .15s;
}
.btn-pdf:hover { background: var(--accent); color: white; }
.transcript-downloads { display: flex; gap: 0.5rem; }

/* Episode navigation */
.episode-nav {
  display: flex; justify-content: space-between;
  padding: 1.5rem 0;
  border-top: 1px solid var(--border);
  margin-top: 2rem;
  gap: 1rem;
}
.episode-nav a {
  font-size: 0.875rem;
  color: var(--text-secondary);
  display: flex; align-items: center; gap: 0.35rem;
  max-width: 45%;
}
.episode-nav a:hover { color: var(--accent); }
.episode-nav .next { margin-left: auto; text-align: right; }

/* ── Breadcrumb ──────────────────────────────────────── */
.breadcrumb {
  font-size: 0.8rem;
  color: var(--text-secondary);
  padding: 1rem 0 0;
}
.breadcrumb a { color: var(--text-secondary); }
.breadcrumb a:hover { color: var(--accent); }
.breadcrumb span { margin: 0 0.4rem; opacity: 0.5; }

/* ── Search Page ─────────────────────────────────────── */
.search-page { padding: 2.5rem 0 4rem; }
.search-page h1 {
  font-family: var(--font-serif);
  font-size: 1.75rem;
  margin-bottom: 1.5rem;
}
.search-input-large {
  width: 100%;
  max-width: 600px;
  padding: 0.75rem 1rem 0.75rem 2.75rem;
  font-size: 1rem;
  font-family: var(--font-sans);
  border: 2px solid var(--border);
  border-radius: var(--radius);
  background: var(--bg-card) url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='20' height='20' viewBox='0 0 24 24' fill='none' stroke='%2378716C' stroke-width='2' stroke-linecap='round' stroke-linejoin='round'%3E%3Ccircle cx='11' cy='11' r='8'/%3E%3Cline x1='21' y1='21' x2='16.65' y2='16.65'/%3E%3C/svg%3E") no-repeat 0.75rem center;
  outline: none;
  transition: border-color .15s, box-shadow .15s;
}
.search-input-large:focus {
  border-color: var(--accent);
  box-shadow: 0 0 0 3px var(--accent-light);
}
.search-info {
  font-size: 0.875rem;
  color: var(--text-secondary);
  margin: 1rem 0;
}
.search-results { list-style: none; margin-top: 1rem; }
.search-result-item {
  border-bottom: 1px solid var(--border-light);
  padding: 1.25rem 0;
}
.search-result-item:last-child { border-bottom: none; }
.search-result-title {
  font-family: var(--font-serif);
  font-size: 1.1rem;
  font-weight: 600;
  margin-bottom: 0.25rem;
}
.search-result-title a { color: var(--text); }
.search-result-title a:hover { color: var(--accent); }
.search-result-meta {
  font-size: 0.8rem;
  color: var(--text-secondary);
  margin-bottom: 0.4rem;
}
.search-result-excerpt {
  font-size: 0.9rem;
  color: var(--text-secondary);
  line-height: 1.6;
}

/* ── Selected Talks ─────────────────────────────────── */
.selected-section { padding: 3rem 0 2rem; border-bottom: 1px solid var(--border); }
.selected-section h2 {
  font-family: var(--font-serif);
  font-size: 1.5rem;
  font-weight: 600;
  margin-bottom: 0.5rem;
  text-align: center;
  color: var(--text);
}
.selected-section .section-subtitle {
  text-align: center;
  color: var(--text-secondary);
  font-size: 0.9rem;
  margin-bottom: 1.5rem;
}
.selected-grid {
  display: grid;
  grid-template-columns: repeat(auto-fill, minmax(340px, 1fr));
  gap: 0.75rem;
}
.selected-item {
  display: grid;
  grid-template-columns: 2rem 1fr auto;
  gap: 0 0.75rem;
  align-items: baseline;
  padding: 0.75rem 1rem;
  background: var(--bg-card);
  border: 1px solid var(--border-light);
  border-radius: var(--radius-sm);
  transition: border-color .15s, box-shadow .15s;
}
.selected-item:hover {
  border-color: var(--accent-light);
  box-shadow: var(--shadow-sm);
}
.selected-rank {
  font-family: var(--font-serif);
  font-size: 0.85rem;
  font-weight: 600;
  color: var(--accent);
  text-align: right;
}
.selected-info { min-width: 0; }
.selected-title {
  font-family: var(--font-serif);
  font-size: 0.95rem;
  font-weight: 600;
  line-height: 1.3;
  white-space: nowrap;
  overflow: hidden;
  text-overflow: ellipsis;
}
.selected-title a { color: var(--text); }
.selected-title a:hover { color: var(--accent); }
.selected-meta {
  font-size: 0.75rem;
  color: var(--text-secondary);
  white-space: nowrap;
  overflow: hidden;
  text-overflow: ellipsis;
}
.selected-score {
  font-size: 0.75rem;
  font-weight: 600;
  color: var(--accent);
  white-space: nowrap;
  padding: 0.15rem 0.5rem;
  background: var(--accent-light);
  border-radius: 100px;
}
.selected-more {
  display: block;
  text-align: center;
  margin-top: 1.25rem;
  font-size: 0.9rem;
  color: var(--accent);
  font-weight: 500;
}
.selected-more:hover { color: var(--accent-hover); }

/* ── Topics Page ────────────────────────────────────── */
.topics-page { padding: 2.5rem 0 4rem; }
.topics-page h1 {
  font-family: var(--font-serif);
  font-size: clamp(1.75rem, 4vw, 2.25rem);
  font-weight: 600;
  margin-bottom: 0.5rem;
}
.topics-page > p {
  color: var(--text-secondary);
  margin-bottom: 2rem;
  line-height: 1.7;
}
.topics-cloud {
  display: flex;
  flex-wrap: wrap;
  gap: 0.4rem;
  justify-content: center;
  padding: 1.5rem;
  background: var(--bg-card);
  border: 1px solid var(--border);
  border-radius: var(--radius);
  margin-bottom: 2rem;
}
.topic-tag {
  display: inline-block;
  padding: 0.25rem 0.65rem;
  border-radius: 100px;
  background: var(--bg-alt);
  border: 1px solid var(--border-light);
  color: var(--text);
  cursor: pointer;
  transition: all .15s;
  white-space: nowrap;
  line-height: 1.4;
}
.topic-tag:hover {
  background: var(--accent-light);
  border-color: var(--accent);
  color: var(--accent);
}
.topic-tag.active {
  background: var(--accent);
  border-color: var(--accent);
  color: white;
}
.topic-tag .count {
  font-size: 0.7em;
  opacity: 0.6;
  margin-left: 0.2em;
}
.topics-filter {
  display: flex; gap: 0.75rem; align-items: center;
  margin-bottom: 1.5rem;
  flex-wrap: wrap;
}
.topics-filter input {
  flex: 1; min-width: 200px;
  padding: 0.5rem 0.75rem;
  border: 1px solid var(--border);
  border-radius: var(--radius-sm);
  font-size: 0.9rem;
  font-family: var(--font-sans);
  outline: none;
  transition: border-color .15s;
}
.topics-filter input:focus { border-color: var(--accent); }
.topics-active-tag {
  display: none;
  align-items: center; gap: 0.5rem;
  padding: 1rem 1.25rem;
  background: var(--accent-light);
  border-radius: var(--radius);
  margin-bottom: 1.5rem;
}
.topics-active-tag.visible { display: flex; }
.topics-active-tag h2 {
  font-family: var(--font-serif);
  font-size: 1.2rem;
  font-weight: 600;
  color: var(--accent);
  margin: 0;
}
.topics-active-tag .close-tag {
  margin-left: auto;
  cursor: pointer;
  color: var(--accent);
  font-size: 1.2rem;
  padding: 0.2rem 0.5rem;
  border-radius: 4px;
}
.topics-active-tag .close-tag:hover { background: rgba(0,0,0,0.05); }
.topics-results { list-style: none; }
.topics-result-item {
  border: 1px solid var(--border-light);
  border-radius: var(--radius-sm);
  padding: 0.85rem 1.25rem;
  margin-bottom: 0.4rem;
  background: var(--bg-card);
  transition: border-color .15s;
  display: grid;
  grid-template-columns: 1fr auto;
  gap: 0.25rem 1rem;
  align-items: start;
}
.topics-result-item:hover {
  border-color: var(--accent-light);
  box-shadow: var(--shadow-sm);
}
.topics-result-title {
  font-family: var(--font-serif);
  font-size: 1rem;
  font-weight: 600;
  line-height: 1.3;
}
.topics-result-title a { color: var(--text); }
.topics-result-title a:hover { color: var(--accent); }
.topics-result-meta {
  font-size: 0.75rem;
  color: var(--text-secondary);
  grid-column: 1;
}
.topics-result-score {
  font-size: 0.75rem;
  font-weight: 600;
  color: var(--accent);
  white-space: nowrap;
}

/* Selected Talks dedicated page */
.selected-page { padding: 2.5rem 0 4rem; }
.selected-page h1 {
  font-family: var(--font-serif);
  font-size: clamp(1.75rem, 4vw, 2.25rem);
  font-weight: 600;
  margin-bottom: 0.5rem;
}
.selected-page > p {
  color: var(--text-secondary);
  margin-bottom: 2rem;
  line-height: 1.7;
}
.selected-full-list { list-style: none; }
.selected-full-item {
  border: 1px solid var(--border-light);
  border-radius: var(--radius-sm);
  padding: 1rem 1.25rem;
  margin-bottom: 0.5rem;
  background: var(--bg-card);
  transition: border-color .15s, box-shadow .15s;
  display: grid;
  grid-template-columns: 2.5rem 1fr auto;
  gap: 0.5rem 1rem;
  align-items: start;
}
.selected-full-item:hover {
  border-color: var(--accent-light);
  box-shadow: var(--shadow-sm);
}
.selected-full-rank {
  font-family: var(--font-serif);
  font-size: 1rem;
  font-weight: 700;
  color: var(--accent);
  text-align: right;
  padding-top: 0.1rem;
}
.selected-full-info { min-width: 0; }
.selected-full-title {
  font-family: var(--font-serif);
  font-size: 1.05rem;
  font-weight: 600;
  line-height: 1.4;
}
.selected-full-title a { color: var(--text); }
.selected-full-title a:hover { color: var(--accent); }
.selected-full-desc {
  font-size: 0.875rem;
  color: var(--text-secondary);
  line-height: 1.6;
  margin-top: 0.2rem;
}
.selected-full-right {
  text-align: right;
  white-space: nowrap;
}
.selected-full-score {
  font-size: 0.85rem;
  font-weight: 700;
  color: var(--accent);
}
.selected-full-duration {
  font-size: 0.75rem;
  color: var(--text-secondary);
  margin-top: 0.2rem;
}
.selected-full-collection {
  font-size: 0.7rem;
  color: var(--text-secondary);
}

/* ── Footer ──────────────────────────────────────────── */
.site-footer {
  background: var(--bg-alt);
  border-top: 1px solid var(--border);
  padding: 2rem 1.5rem;
  text-align: center;
  font-size: 0.8rem;
  color: var(--text-secondary);
}
.site-footer a { color: var(--text-secondary); text-decoration: underline; }
.site-footer a:hover { color: var(--accent); }

/* ── Responsive ──────────────────────────────────────── */
@media (max-width: 768px) {
  .header-inner { height: 48px; }
  .header-nav { gap: 1rem; }
  .header-search input { width: 120px; }
  .hero { padding: 2.5rem 1.25rem; }
  .hero-stats { gap: 1.5rem; }
  .feeds-grid { grid-template-columns: 1fr; }
  .episode-item { grid-template-columns: 1fr; }
  .episode-duration { order: -1; }
  .episode-nav { flex-direction: column; }
  .episode-nav a { max-width: 100%; }
}
@media (max-width: 480px) {
  .header-search { display: none; }
  .hero h1 { font-size: 1.75rem; }
  .feed-card-header { padding: 1.25rem 1.25rem 0.75rem; }
  .feed-card-body { padding: 0 1.25rem 1.25rem; }
}
"""

CSS_HASH = hashlib.md5(CSS.encode()).hexdigest()[:8]

# ── Icônes SVG ─────────────────────────────────────────

SVG_SEARCH = '<svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="11" cy="11" r="8"/><line x1="21" y1="21" x2="16.65" y2="16.65"/></svg>'
SVG_ARROW_LEFT = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="19" y1="12" x2="5" y2="12"/><polyline points="12 19 5 12 12 5"/></svg>'
SVG_ARROW_RIGHT = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="5" y1="12" x2="19" y2="12"/><polyline points="12 5 19 12 12 19"/></svg>'
SVG_SPEAKER = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M20 21v-2a4 4 0 0 0-4-4H8a4 4 0 0 0-4 4v2"/><circle cx="12" cy="7" r="4"/></svg>'
SVG_CLOCK = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg>'
SVG_FOLDER = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M22 19a2 2 0 0 1-2 2H4a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h5l2 3h9a2 2 0 0 1 2 2z"/></svg>'
SVG_EPISODES = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="8" y1="6" x2="21" y2="6"/><line x1="8" y1="12" x2="21" y2="12"/><line x1="8" y1="18" x2="21" y2="18"/><line x1="3" y1="6" x2="3.01" y2="6"/><line x1="3" y1="12" x2="3.01" y2="12"/><line x1="3" y1="18" x2="3.01" y2="18"/></svg>'
SVG_BOOK = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M4 19.5A2.5 2.5 0 0 1 6.5 17H20"/><path d="M6.5 2H20v20H6.5A2.5 2.5 0 0 1 4 19.5v-15A2.5 2.5 0 0 1 6.5 2z"/></svg>'
SVG_DOWNLOAD = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>'
SVG_HEADPHONES = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M3 18v-6a9 9 0 0 1 18 0v6"/><path d="M21 19a2 2 0 0 1-2 2h-1a2 2 0 0 1-2-2v-3a2 2 0 0 1 2-2h3zM3 19a2 2 0 0 0 2 2h1a2 2 0 0 0 2-2v-3a2 2 0 0 0-2-2H3z"/></svg>'


# ── Templates HTML ─────────────────────────────────────

def html_base(title, body_html, breadcrumbs=None, extra_head="", body_class=""):
    """Enveloppe le contenu dans le template de base HTML."""
    bc = ""
    if breadcrumbs:
        items = [f'<a href="{url}">{label}</a>' for label, url in breadcrumbs]
        bc = f'<div class="container"><nav class="breadcrumb">{"<span>›</span>".join(items)}</nav></div>'

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{h(title)}</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Crimson+Pro:ital,wght@0,400;0,600;0,700;1,400&family=Inter:wght@400;500;600&display=swap" rel="stylesheet">
  <link rel="icon" type="image/png" href="{base('/favicon.png')}">
  <link rel="apple-touch-icon" href="{base('/apple-touch-icon.png')}">
  <link rel="stylesheet" href="{base('/style.css')}?v={CSS_HASH}">
  {extra_head}
</head>
<body class="{body_class}">
  <header class="site-header">
    <div class="header-inner">
      <a href="{base('/')}" class="site-logo">Satipanya</a>
      <nav class="header-nav">
        <a href="{base('/')}">Collections</a>
        <a href="{base('/selected/')}">Selected</a>
        <a href="{base('/topics/')}">Topics</a>
        <a href="{base('/search.html')}">Search</a>
        <form class="header-search" action="{base('/search.html')}" method="get">
          {SVG_SEARCH}
          <input type="search" name="q" placeholder="Search talks…" autocomplete="off">
        </form>
      </nav>
    </div>
  </header>
  {bc}
  <main>
{body_html}
  </main>
  <footer class="site-footer">
    <p>Talks by <a href="https://www.satipanya.org.uk" target="_blank" rel="noopener">Satipanya Buddhist Retreat</a>.
    Site compiled from the original audio archive.</p>
    <p>Transcriptions produced locally using Swiss low-carbon electricity. Corrections and rewriting by cloud-hosted AI.
    <br>Source code and full archive on <a href="https://github.com/funkypitt/satipanya-dharma-library">GitHub</a>.</p>
  </footer>
</body>
</html>"""


# ── Générateurs de pages ───────────────────────────────

def build_homepage(catalog, selected_talks=None):
    """Génère la page d'accueil."""
    # Calcul stats globales
    total_episodes = 0
    total_hours = 0
    total_articles = 0
    total_essays = 0
    for slug in FEED_ORDER:
        fdata = catalog.get(slug)
        if not fdata:
            continue
        text_feed = is_text_feed(fdata)
        local_audio = fdata.get("source_type") == "local_audio"
        for season in fdata.get("seasons", []):
            for ep in season.get("episodes", []):
                stem = ep_stem(ep)
                if not stem:
                    continue
                if local_audio:
                    # Retreat talks: count hours from audio, transcripts from articles
                    dur = ep.get("duration_seconds", 0)
                    if dur > 0:
                        total_hours += dur / 3600
                    if stem and (ARTICLES_DIR / slug / f"{stem}.txt").exists():
                        total_articles += 1
                elif text_feed:
                    if ep.get("word_count", 0) > 0:
                        total_essays += 1
                else:
                    dur = ep.get("duration_seconds", 0)
                    if dur == 0:
                        continue
                    total_episodes += 1
                    total_hours += dur / 3600
                    if stem and (ARTICLES_DIR / slug / f"{stem}.txt").exists():
                        total_articles += 1

    # Audio feed cards
    audio_cards = []
    for slug in AUDIO_FEED_ORDER:
        fdata = catalog.get(slug)
        if not fdata:
            continue
        emoji = FEED_EMOJI.get(slug, "📖")
        ep_count = 0
        for season in fdata.get("seasons", []):
            for ep in season.get("episodes", []):
                if ep.get("duration_seconds", 0) > 0:
                    ep_count += 1
        # Subscribe link: YouTube channel or podcast://
        if slug == "youtube-talks":
            subscribe_url = "https://www.youtube.com/@satipanya-insight"
            subscribe_label = f'{SVG_HEADPHONES} YouTube'
        else:
            subscribe_url = f"podcast://{FEEDS_BASE_URL.replace('https://', '')}/{slug}.xml"
            subscribe_label = f'{SVG_HEADPHONES} Subscribe'
        # Liens livres (toujours inclus — générés par build_books.py)
        book_links = (
            f'<a href="{base(f"/books/{slug}.pdf")}" class="subscribe">{SVG_BOOK} PDF</a>'
            f'<a href="{base(f"/books/{slug}.epub")}" class="subscribe">{SVG_BOOK} EPUB</a>'
            f'<a href="{base(f"/books/{slug}.docx")}" class="subscribe">{SVG_BOOK} DOCX</a>'
        )
        audio_cards.append(f"""
      <div class="feed-card">
        <div class="feed-card-header">
          <div class="feed-card-emoji">{emoji}</div>
          <h3><a href="{base(f"/{slug}/")}">{h(fdata['name'].replace('Satipanya — ', ''))}</a></h3>
        </div>
        <div class="feed-card-body">
          <p>{h(fdata.get('description', ''))}</p>
          <div class="feed-card-meta">
            <span>{SVG_EPISODES} {ep_count} talks</span>
            <a href="{subscribe_url}" class="subscribe">{subscribe_label}</a>
            {book_links}
          </div>
        </div>
      </div>""")

    # Text feed cards
    text_cards = []
    for slug in TEXT_FEED_ORDER:
        fdata = catalog.get(slug)
        if not fdata:
            continue
        emoji = FEED_EMOJI.get(slug, "📖")
        ep_count = 0
        for season in fdata.get("seasons", []):
            for ep in season.get("episodes", []):
                if ep.get("word_count", 0) > 0 or ep_stem(ep):
                    ep_count += 1
        book_links = (
            f'<a href="{base(f"/books/{slug}.pdf")}" class="subscribe">{SVG_BOOK} PDF</a>'
            f'<a href="{base(f"/books/{slug}.epub")}" class="subscribe">{SVG_BOOK} EPUB</a>'
            f'<a href="{base(f"/books/{slug}.docx")}" class="subscribe">{SVG_BOOK} DOCX</a>'
        )
        text_cards.append(f"""
      <div class="feed-card">
        <div class="feed-card-header">
          <div class="feed-card-emoji">{emoji}</div>
          <h3><a href="{base(f"/{slug}/")}">{h(fdata['name'].replace('Satipanya — ', ''))}</a></h3>
        </div>
        <div class="feed-card-body">
          <p>{h(fdata.get('description', ''))}</p>
          <div class="feed-card-meta">
            <span>{SVG_EPISODES} {ep_count} {"transcripts" if fdata.get("source_type") == "local_audio" else "essays"}</span>
            {book_links}
          </div>
        </div>
      </div>""")

    selected_card = build_selected_homepage_card(selected_talks or [])

    # Hero stats: include essays if any
    essays_stat = f'<div><strong>{total_essays}</strong> essays</div>' if total_essays > 0 else ''

    # Text section: only show if there are text collections
    text_section = ""
    if text_cards:
        text_section = f"""
    <section class="feeds-section">
      <div class="container">
        <h2>Read</h2>
        <div class="feeds-grid">
          {"".join(text_cards)}
        </div>
      </div>
    </section>"""

    body = f"""
    <section class="hero">
      <div class="container">
        <h1>{SITE_TITLE}</h1>
        <p>{SITE_TAGLINE}</p>
        <div class="hero-stats">
          <div><strong>{total_episodes}</strong> talks</div>
          <div><strong>{int(total_hours)}</strong> hours</div>
          <div><strong>{total_articles}</strong> transcripts</div>
          {essays_stat}
        </div>
      </div>
    </section>
    <section class="feeds-section">
      <div class="container">
        <h2>Listen</h2>
        <div class="feeds-grid">
          {selected_card}
          {"".join(audio_cards)}
        </div>
      </div>
    </section>
    {text_section}"""

    return html_base(SITE_TITLE, body, body_class="page-home")


def build_feed_page(slug, fdata, catalog):
    """Génère la page d'une collection/feed."""
    name = fdata["name"].replace("Satipanya — ", "")
    desc = fdata.get("description", "")
    text_feed = is_text_feed(fdata)

    newest_first = slug in FEEDS_NEWEST_FIRST
    seasons_list = list(fdata.get("seasons", []))
    if newest_first:
        seasons_list = sorted(seasons_list, key=_season_sort_key, reverse=True)

    seasons_html = []
    for season in seasons_list:
        episodes = list(season.get("episodes", []))
        if newest_first:
            episodes = list(reversed(episodes))
        items = []
        for ep in episodes:
            stem = ep_stem(ep)
            if not stem:
                continue
            if text_feed:
                if ep.get("word_count", 0) == 0:
                    continue
                time_str = format_reading_time(ep.get("reading_minutes", 0))
            else:
                dur = ep.get("duration_seconds", 0)
                if dur == 0:
                    continue
                time_str = format_duration(dur)
            meta = load_metadata(slug, stem)
            title = meta.get("title_clean", ep.get("title", "Untitled"))
            desc_short = meta.get("description_short", ep.get("description_short", ""))
            url = ep_url(slug, stem)
            items.append(f"""
          <li class="episode-item">
            <div class="episode-title"><a href="{url}">{h(title)}</a></div>
            <div class="episode-duration">{SVG_CLOCK} {time_str}</div>
            <div class="episode-desc">{h(desc_short)}</div>
          </li>""")

        if items:
            s_name = season.get("name", f"Season {season.get('number', '?')}")
            seasons_html.append(f"""
        <section class="season-section">
          <h2 class="season-title">{h(s_name)}</h2>
          <ul class="episode-list">{"".join(items)}</ul>
        </section>""")

    body = f"""
    <div class="container">
      <div class="feed-header">
        <h1>{FEED_EMOJI.get(slug, '📖')} {h(name)}</h1>
        <p>{h(desc)}</p>
      </div>
      {"".join(seasons_html)}
    </div>"""

    breadcrumbs = [("Home", base("/")), (name, base(f"/{slug}/"))]
    return html_base(f"{name} — {SITE_TITLE}", body, breadcrumbs=breadcrumbs)


def build_episode_page(slug, ep, prev_ep, next_ep, feed_name):
    """Génère la page d'un épisode."""
    stem = ep_stem(ep)
    meta = load_metadata(slug, stem) if stem else {}
    article = load_article(slug, stem) if stem else None
    text_ep = is_text_episode(ep)

    title = meta.get("title_clean", ep.get("title", "Untitled"))
    desc_long = meta.get("description_long") or ep.get("description_long", "")
    desc_short = meta.get("description_short") or ep.get("description_short", "")
    keywords = meta.get("keywords", [])
    speaker = ep.get("speaker", "")
    dur = ep.get("duration_seconds", 0)
    audio_url = ep.get("url", "")
    clean_feed = feed_name.replace("Satipanya — ", "")

    # Meta bar: reading time for text, duration for audio
    if text_ep:
        reading_min = ep.get("reading_minutes", 0)
        word_count = ep.get("word_count", 0)
        time_html = f'<span>{SVG_CLOCK} {format_reading_time(reading_min)}</span>'
        if word_count:
            time_html += f' <span>({word_count:,} words)</span>'
        # Show original talk duration for local_audio transcripts
        if ep.get("source_type") == "local_audio" and dur:
            time_html += f' <span>· Original talk: {format_duration(dur)}</span>'
    else:
        time_html = f'<span>{SVG_CLOCK} {format_duration(dur)}</span>'

    # Description HTML
    desc_html = ""
    if desc_long:
        # Retirer la note auto-générée
        desc_text = desc_long.replace("(This description was generated automatically, inaccuracies may happen in the process.)", "").strip()
        desc_paras = [f"<p>{h(p.strip())}</p>" for p in desc_text.split("\n\n") if p.strip()]
        desc_html = f'<div class="episode-description">{"".join(desc_paras)}</div>'
    elif desc_short:
        desc_html = f'<div class="episode-description"><p>{h(desc_short)}</p></div>'

    # Audio / Video — only for non-text episodes
    audio_html = ""
    if not text_ep and audio_url:
        # YouTube: embed iframe; otherwise: audio player
        yt_match = re.search(r'youtube\.com/watch\?v=([A-Za-z0-9_-]+)', audio_url)
        if yt_match:
            yt_id = yt_match.group(1)
            audio_html = f"""
      <div class="audio-player">
        <div class="youtube-embed">
          <iframe src="https://www.youtube.com/embed/{yt_id}" frameborder="0"
                  allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture"
                  allowfullscreen></iframe>
        </div>
        <div class="audio-source">Source: <a href="{h(audio_url)}" target="_blank">YouTube</a></div>
      </div>"""
        else:
            audio_html = f"""
      <div class="audio-player">
        <audio controls preload="none" src="{h(audio_url)}">
          Your browser does not support the audio element.
        </audio>
        <div class="audio-source">Source: <a href="{h(audio_url)}" target="_blank">satipanya.org.uk</a></div>
      </div>"""

    # Source link for text episodes (skip for local_audio — no external URL)
    source_html = ""
    if text_ep and audio_url and not audio_url.startswith("local://"):
        source_html = f'<p class="original-source">Original source: <a href="{h(audio_url)}" target="_blank">satipanya.org.uk</a></p>'

    # Transcript / Full Text
    transcript_html = ""
    download_links = ""
    if article and stem:
        pdf_name = f"{stem}.pdf"
        epub_name = f"{stem}.epub"
        docx_name = f"{stem}.docx"
        download_links = (
            f'<div class="transcript-downloads">'
            f'<a href="{base(f"/{slug}/{pdf_name}")}" class="btn-pdf" download>{SVG_DOWNLOAD} PDF</a>'
            f'<a href="{base(f"/{slug}/{epub_name}")}" class="btn-pdf" download>{SVG_BOOK} EPUB</a>'
            f'<a href="{base(f"/{slug}/{docx_name}")}" class="btn-pdf" download>{SVG_BOOK} DOCX</a>'
            f'</div>'
        )
        section_title = "Full Text" if text_ep else "Transcript"
        transcript_html = f"""
      <section class="transcript-section">
        <div class="transcript-header">
          <h2>{section_title}</h2>
          {download_links}
        </div>
        <div class="transcript-text">
          {article_to_html(article)}
        </div>
      </section>"""
    elif not text_ep:
        transcript_html = """
      <section class="transcript-section">
        <h2>Transcript</h2>
        <p class="transcript-unavailable">Written transcript not yet available for this talk.</p>
      </section>"""

    # Navigation prev/next
    nav_html = ""
    nav_parts = []
    if prev_ep:
        p_stem = ep_stem(prev_ep)
        p_meta = load_metadata(slug, p_stem) if p_stem else {}
        p_title = p_meta.get("title_clean", prev_ep.get("title", "Previous"))
        nav_parts.append(f'<a href="{ep_url(slug, p_stem)}" class="prev">{SVG_ARROW_LEFT} {h(p_title)}</a>')
    else:
        nav_parts.append('<span></span>')
    if next_ep:
        n_stem = ep_stem(next_ep)
        n_meta = load_metadata(slug, n_stem) if n_stem else {}
        n_title = n_meta.get("title_clean", next_ep.get("title", "Next"))
        nav_parts.append(f'<a href="{ep_url(slug, n_stem)}" class="next">{h(n_title)} {SVG_ARROW_RIGHT}</a>')
    if any(s.strip() for s in nav_parts):
        nav_html = f'<nav class="episode-nav">{"".join(nav_parts)}</nav>'

    body = f"""
    <div class="container">
      <div class="episode-header">
        <h1>{h(title)}</h1>
        <div class="episode-meta">
          <span>{SVG_SPEAKER} {h(speaker)}</span>
          {time_html}
          <span>{SVG_FOLDER} <a href="{base(f"/{slug}/")}">{h(clean_feed)}</a></span>
        </div>
      </div>
      {audio_html}
      {source_html}
      {desc_html}
      {keyword_tags_html(keywords)}
      {transcript_html}
      {nav_html}
    </div>"""

    breadcrumbs = [("Home", base("/")), (clean_feed, base(f"/{slug}/")), (title, base(f"/{slug}/{stem}.html"))]
    return html_base(f"{title} — {SITE_TITLE}", body, breadcrumbs=breadcrumbs)


def build_search_page():
    """Génère la page de recherche."""
    body = """
    <div class="container search-page">
      <h1>Search the Dharma Library</h1>
      <input type="search" id="search-input" class="search-input-large"
             placeholder="Search by topic, Pali term, speaker, keyword…" autofocus>
      <div id="search-info" class="search-info"></div>
      <ul id="search-results" class="search-results"></ul>
    </div>"""

    search_js = """
<script src="https://cdn.jsdelivr.net/npm/minisearch@7.1.1/dist/umd/index.min.js"></script>
<script>
(function() {
  let ms, docs;
  const input = document.getElementById('search-input');
  const resultsEl = document.getElementById('search-results');
  const infoEl = document.getElementById('search-info');

  // Charger l'index
  fetch('__BASE__/search-index.json')
    .then(r => r.json())
    .then(data => {
      docs = data;
      ms = new MiniSearch({
        fields: ['title', 'description', 'keywords', 'speaker', 'transcript_excerpt'],
        storeFields: ['title', 'feed', 'speaker', 'duration', 'description_short', 'url'],
        searchOptions: {
          boost: { title: 3, keywords: 2, description: 1.5 },
          fuzzy: 0.2,
          prefix: true
        }
      });
      ms.addAll(docs);
      infoEl.textContent = docs.length + ' talks indexed';
      // Chercher si query dans l'URL
      const q = new URLSearchParams(location.search).get('q');
      if (q) { input.value = q; doSearch(q); }
    });

  input.addEventListener('input', (e) => doSearch(e.target.value));

  function doSearch(query) {
    if (!ms || !query.trim()) {
      resultsEl.innerHTML = '';
      if (ms) infoEl.textContent = docs.length + ' talks indexed';
      return;
    }
    const results = ms.search(query, { combineWith: 'AND' });
    infoEl.textContent = results.length + ' result' + (results.length !== 1 ? 's' : '');
    resultsEl.innerHTML = results.slice(0, 50).map(r => `
      <li class="search-result-item">
        <div class="search-result-title"><a href="${r.url}">${esc(r.title)}</a></div>
        <div class="search-result-meta">${esc(r.feed)} · ${esc(r.speaker)} · ${r.duration}</div>
        <div class="search-result-excerpt">${esc(r.description_short)}</div>
      </li>
    `).join('');
  }

  function esc(s) {
    if (!s) return '';
    const d = document.createElement('div');
    d.textContent = s;
    return d.innerHTML;
  }
})();
</script>"""

    page = html_base("Search — " + SITE_TITLE, body, extra_head="", body_class="page-search")
    search_js_resolved = search_js.replace("__BASE__", SITE_BASE_PATH)
    return page.replace("</body>", search_js_resolved + "\n</body>")


def build_search_index(catalog):
    """Construit l'index JSON pour la recherche côté client."""
    items = []
    idx = 0
    for slug in FEED_ORDER:
        fdata = catalog.get(slug)
        if not fdata:
            continue
        feed_name = fdata["name"].replace("Satipanya — ", "")
        text_feed = is_text_feed(fdata)
        for season in fdata.get("seasons", []):
            for ep in season.get("episodes", []):
                stem = ep_stem(ep)
                if not stem:
                    continue
                if text_feed:
                    if ep.get("word_count", 0) == 0:
                        continue
                    time_str = format_reading_time(ep.get("reading_minutes", 0))
                else:
                    dur = ep.get("duration_seconds", 0)
                    if dur == 0:
                        continue
                    time_str = format_duration(dur)
                meta = load_metadata(slug, stem)
                article = load_article(slug, stem)

                title = meta.get("title_clean", ep.get("title", ""))
                desc_short = meta.get("description_short") or ep.get("description_short", "")
                desc_long = meta.get("description_long") or ep.get("description_long", "")
                # Retirer la note auto-générée
                desc_long = desc_long.replace(
                    "(This description was generated automatically, inaccuracies may happen in the process.)", ""
                ).strip()
                keywords = " ".join(meta.get("keywords", []))

                # Extrait du transcript pour la recherche
                transcript_excerpt = ""
                if article:
                    transcript_excerpt = article[:2000]

                items.append({
                    "id": idx,
                    "title": title,
                    "feed": feed_name,
                    "speaker": ep.get("speaker", ""),
                    "duration": time_str,
                    "description": f"{desc_short} {desc_long}",
                    "description_short": desc_short,
                    "keywords": keywords,
                    "transcript_excerpt": transcript_excerpt,
                    "url": ep_url(slug, stem),
                })
                idx += 1
    return items


# ── Selected Talks (top par qualité littéraire) ────────

SELECTED_TALKS_COUNT = 120   # Nombre de talks dans la sélection
SELECTED_HOMEPAGE_COUNT = 12  # Nombre de talks affichés sur la page d'accueil

SVG_STAR = '<svg width="14" height="14" viewBox="0 0 24 24" fill="currentColor" stroke="none"><path d="M12 2l3.09 6.26L22 9.27l-5 4.87 1.18 6.88L12 17.77l-6.18 3.25L7 14.14 2 9.27l6.91-1.01L12 2z"/></svg>'


def collect_selected_talks(catalog, count=SELECTED_TALKS_COUNT):
    """Collecte les top N épisodes par lite_score à travers toutes les collections.

    S'assure que deux talks de Noirin ne sont jamais adjacents, tout en
    respectant l'ordre par score décroissant autant que possible.
    """
    scored = []
    for key, fdata in catalog.items():
        slug = fdata.get("slug", key.replace("_", "-"))
        if slug not in FEED_ORDER:
            continue
        feed_name = fdata["name"].replace("Satipanya — ", "")
        text_feed = is_text_feed(fdata)
        for season in fdata.get("seasons", []):
            for ep in season.get("episodes", []):
                score = ep.get("lite_score")
                if score is None or score < 1:
                    continue
                stem = ep_stem(ep)
                if not stem:
                    continue
                if text_feed:
                    if ep.get("word_count", 0) == 0:
                        continue
                else:
                    if ep.get("duration_seconds", 0) == 0:
                        continue
                meta = load_metadata(slug, stem)
                title = meta.get("title_clean", ep.get("title", "Untitled"))
                desc_short = meta.get("description_short") or ep.get("description_short", "")
                if text_feed:
                    dur_display = ep.get("reading_minutes", 0)
                else:
                    dur_display = ep.get("duration_seconds", 0)
                scored.append({
                    "title": title,
                    "slug": slug,
                    "feed_name": feed_name,
                    "stem": stem,
                    "score": score,
                    "duration": dur_display,
                    "speaker": ep.get("speaker", ""),
                    "desc_short": desc_short,
                    "url": ep_url(slug, stem),
                    "is_text": text_feed,
                })
    # Tri par score décroissant, puis par titre pour départager
    scored.sort(key=lambda x: (-x["score"], x["title"]))

    # S'assurer que deux talks de Noirin ne soient jamais adjacents
    noirin = [t for t in scored if "noirin" in t["speaker"].lower()]
    bhante = [t for t in scored if "noirin" not in t["speaker"].lower()]
    result = []
    ni, bi = 0, 0
    while ni < len(noirin) or bi < len(bhante):
        prev_noirin = result and "noirin" in result[-1]["speaker"].lower()
        if prev_noirin:
            # Dernier était Noirin → on doit placer un Bhante
            if bi < len(bhante):
                result.append(bhante[bi]); bi += 1
            elif ni < len(noirin):
                result.append(noirin[ni]); ni += 1  # pas le choix
        else:
            # Dernier était Bhante (ou liste vide) → meilleur score entre les deux
            n_score = noirin[ni]["score"] if ni < len(noirin) else -1
            b_score = bhante[bi]["score"] if bi < len(bhante) else -1
            if n_score >= b_score and ni < len(noirin):
                result.append(noirin[ni]); ni += 1
            elif bi < len(bhante):
                result.append(bhante[bi]); bi += 1
            elif ni < len(noirin):
                result.append(noirin[ni]); ni += 1
    scored = result

    return scored[:count]


def build_selected_homepage_card(selected_talks):
    """Génère une feed-card pour la sélection, à intégrer dans la grille de collections."""
    if not selected_talks:
        return ""

    book_links = (
        f'<a href="{base("/books/selected-talks.pdf")}" class="subscribe">{SVG_BOOK} PDF</a>'
        f'<a href="{base("/books/selected-talks.epub")}" class="subscribe">{SVG_BOOK} EPUB</a>'
        f'<a href="{base("/books/selected-talks.docx")}" class="subscribe">{SVG_BOOK} DOCX</a>'
    )

    return f"""
      <div class="feed-card">
        <div class="feed-card-header">
          <div class="feed-card-emoji">{SVG_STAR}</div>
          <h3><a href="{base('/selected/')}">Selected Talks</a></h3>
        </div>
        <div class="feed-card-body">
          <p>A selection of {len(selected_talks)} interesting talks if you don't know where to start.</p>
          <div class="feed-card-meta">
            <span>{SVG_EPISODES} {len(selected_talks)} talks</span>
            {book_links}
          </div>
        </div>
      </div>"""


def build_selected_page(selected_talks):
    """Génère la page dédiée avec les top talks."""
    if not selected_talks:
        return None

    items = []
    for rank, talk in enumerate(selected_talks, 1):
        if talk.get("is_text"):
            dur_str = format_reading_time(talk["duration"])
        else:
            dur_str = format_duration(talk["duration"])
        items.append(f"""
        <li class="selected-full-item">
          <div class="selected-full-rank">{rank}</div>
          <div class="selected-full-info">
            <div class="selected-full-title"><a href="{talk['url']}">{h(talk['title'])}</a></div>
            <div class="selected-full-desc">{h(talk['desc_short'][:150])}</div>
          </div>
          <div class="selected-full-right">
            <div class="selected-full-duration">{SVG_CLOCK} {dur_str}</div>
            <div class="selected-full-collection">{h(talk['feed_name'])}</div>
          </div>
        </li>""")

    body = f"""
    <div class="container selected-page">
      <h1>{SVG_STAR} Selected Talks</h1>
      <p>A selection of {len(selected_talks)} interesting talks from across all collections.
         If you don't know where to start, these are a good place to begin.</p>
      <ul class="selected-full-list">{"".join(items)}</ul>
    </div>"""

    breadcrumbs = [("Home", base("/")), ("Selected Talks", base("/selected/"))]
    return html_base(f"Selected Talks — {SITE_TITLE}", body, breadcrumbs=breadcrumbs)


# ── Topics (tag cloud + filtered episodes) ─────────────

MIN_TAG_COUNT = 2  # Tags apparaissant au moins N fois


def build_topics_index(catalog):
    """Construit l'index JSON {tag → [episodes]} pour la page Topics.

    Chaque épisode dans la liste est trié par lite_score décroissant.
    """
    from collections import defaultdict
    tag_episodes = defaultdict(list)

    for key, fdata in catalog.items():
        slug = fdata.get("slug", key.replace("_", "-"))
        if slug not in FEED_ORDER:
            continue
        feed_name = fdata["name"].replace("Satipanya — ", "")
        text_feed = is_text_feed(fdata)
        for season in fdata.get("seasons", []):
            for ep in season.get("episodes", []):
                stem = ep_stem(ep)
                if not stem:
                    continue
                if text_feed:
                    if ep.get("word_count", 0) == 0:
                        continue
                    time_str = format_reading_time(ep.get("reading_minutes", 0))
                else:
                    if ep.get("duration_seconds", 0) == 0:
                        continue
                    time_str = format_duration(ep.get("duration_seconds", 0))
                meta = load_metadata(slug, stem)
                keywords = meta.get("keywords", [])
                if not keywords:
                    continue

                title = meta.get("title_clean", ep.get("title", "Untitled"))
                score = ep.get("lite_score", 0)
                ep_info = {
                    "t": title,
                    "u": ep_url(slug, stem),
                    "f": feed_name,
                    "d": time_str,
                    "s": score,
                }

                for kw in keywords:
                    tag_episodes[kw].append(ep_info)

    # Filtrer les tags trop rares et trier les épisodes par score
    result = {}
    for tag, eps in sorted(tag_episodes.items()):
        if len(eps) < MIN_TAG_COUNT:
            continue
        eps.sort(key=lambda x: -x["s"])
        result[tag] = eps

    return result


def build_topics_page(topics_index):
    """Génère la page Topics avec tag cloud et résultats filtrés."""
    if not topics_index:
        return None

    # Calculer les tailles de font pour le cloud
    counts = {tag: len(eps) for tag, eps in topics_index.items()}
    max_count = max(counts.values()) if counts else 1
    min_count = min(counts.values()) if counts else 1

    # Trier les tags par fréquence décroissante pour le cloud
    sorted_tags = sorted(counts.items(), key=lambda x: (-x[1], x[0]))

    cloud_items = []
    for tag, count in sorted_tags:
        # Font size: 0.7rem (rare) → 1.6rem (très fréquent)
        if max_count > min_count:
            ratio = (count - min_count) / (max_count - min_count)
        else:
            ratio = 0.5
        size = 0.7 + ratio * 0.9
        cloud_items.append(
            f'<span class="topic-tag" data-tag="{h(tag)}" style="font-size:{size:.2f}rem">'
            f'{h(tag)}<span class="count">{count}</span></span>'
        )

    body = f"""
    <div class="container topics-page">
      <h1>Explore by Topic</h1>
      <p>{len(topics_index)} topics across {sum(counts.values())} tag occurrences.
         Click a topic to see related talks.</p>
      <div class="topics-filter">
        <input type="text" id="topic-filter" placeholder="Filter topics…">
      </div>
      <div class="topics-cloud" id="topics-cloud">
        {"".join(cloud_items)}
      </div>
      <div class="topics-active-tag" id="active-tag">
        <h2 id="active-tag-name"></h2>
        <span id="active-tag-count" style="font-size:0.85rem;color:var(--text-secondary)"></span>
        <span class="close-tag" id="close-tag">✕</span>
      </div>
      <ul class="topics-results" id="topics-results"></ul>
    </div>"""

    topics_js = """
<script>
(function() {
  let topicsData = null;
  const cloud = document.getElementById('topics-cloud');
  const resultsEl = document.getElementById('topics-results');
  const activeTag = document.getElementById('active-tag');
  const activeTagName = document.getElementById('active-tag-name');
  const activeTagCount = document.getElementById('active-tag-count');
  const closeTag = document.getElementById('close-tag');
  const filterInput = document.getElementById('topic-filter');
  const allTags = cloud.querySelectorAll('.topic-tag');

  fetch('__BASE__/topics-index.json')
    .then(r => r.json())
    .then(data => { topicsData = data; checkUrlTag(); });

  // Filter topics in cloud
  filterInput.addEventListener('input', function() {
    const q = this.value.toLowerCase();
    allTags.forEach(tag => {
      tag.style.display = tag.dataset.tag.toLowerCase().includes(q) ? '' : 'none';
    });
  });

  // Click on tag
  cloud.addEventListener('click', function(e) {
    const tag = e.target.closest('.topic-tag');
    if (tag) showTag(tag.dataset.tag);
  });

  closeTag.addEventListener('click', function() {
    activeTag.classList.remove('visible');
    resultsEl.innerHTML = '';
    allTags.forEach(t => t.classList.remove('active'));
    history.replaceState(null, '', location.pathname);
  });

  function checkUrlTag() {
    const t = new URLSearchParams(location.search).get('t');
    if (t && topicsData && topicsData[t]) showTag(t);
  }

  function showTag(tagName) {
    if (!topicsData || !topicsData[tagName]) return;
    const eps = topicsData[tagName];

    allTags.forEach(t => {
      t.classList.toggle('active', t.dataset.tag === tagName);
    });

    activeTagName.textContent = tagName;
    activeTagCount.textContent = eps.length + ' talk' + (eps.length !== 1 ? 's' : '');
    activeTag.classList.add('visible');

    resultsEl.innerHTML = eps.map((ep, i) => {
      const scoreHtml = ep.s > 0
        ? '<span class="topics-result-score">★ ' + ep.s + '</span>'
        : '';
      return '<li class="topics-result-item">' +
        '<div class="topics-result-title"><a href="' + ep.u + '">' + esc(ep.t) + '</a></div>' +
        scoreHtml +
        '<div class="topics-result-meta">' + esc(ep.f) + ' · ' + ep.d + '</div>' +
        '</li>';
    }).join('');

    history.replaceState(null, '', location.pathname + '?t=' + encodeURIComponent(tagName));
  }

  function esc(s) {
    if (!s) return '';
    const d = document.createElement('div');
    d.textContent = s;
    return d.innerHTML;
  }
})();
</script>"""

    breadcrumbs = [("Home", base("/")), ("Topics", base("/topics/"))]
    page = html_base(f"Topics — {SITE_TITLE}", body, breadcrumbs=breadcrumbs)
    topics_js_resolved = topics_js.replace("__BASE__", SITE_BASE_PATH)
    return page.replace("</body>", topics_js_resolved + "\n</body>")


# ── Construction du site ───────────────────────────────

def build_site():
    """Point d'entrée : génère tout le site statique."""
    print("=" * 60)
    print("Building Satipanya Dharma Library static site")
    print("=" * 60)

    # Charger le catalogue
    with open(CATALOG_PATH) as f:
        raw_catalog = json.load(f)

    # Réindexer par slug
    catalog = {}
    for key, fdata in raw_catalog.items():
        slug = fdata.get("slug", key.replace("_", "-"))
        catalog[slug] = fdata

    # Nettoyer et créer le répertoire de sortie (préserver books/)
    if SITE_DIR.exists():
        for item in SITE_DIR.iterdir():
            if item.name == "books":
                continue
            if item.is_dir():
                shutil.rmtree(item)
            else:
                item.unlink()
    else:
        SITE_DIR.mkdir(parents=True)

    # CSS
    (SITE_DIR / "style.css").write_text(CSS, encoding="utf-8")
    print(f"  ✓ style.css")

    # Covers
    covers_out = SITE_DIR / "covers"
    covers_out.mkdir()
    for png in COVERS_DIR.glob("*.png"):
        shutil.copy2(png, covers_out / png.name)
    print(f"  ✓ covers/ ({len(list(covers_out.glob('*.png')))} images)")

    # Favicon (depuis le logo Satipanya)
    favicon_src = COVERS_DIR / "favicon-source.png"
    if favicon_src.exists():
        import subprocess
        subprocess.run(["convert", str(favicon_src), "-resize", "32x32",
                        str(SITE_DIR / "favicon.png")], check=True)
        subprocess.run(["convert", str(favicon_src), "-resize", "180x180",
                        str(SITE_DIR / "apple-touch-icon.png")], check=True)
        print(f"  ✓ favicon.png + apple-touch-icon.png")

    # RSS feeds (copiés depuis feeds/ pour tout avoir dans site/)
    feeds_src = PROJECT_DIR / "feeds"
    if feeds_src.is_dir():
        n_feeds = 0
        for xml in feeds_src.glob("*.xml"):
            shutil.copy2(xml, SITE_DIR / xml.name)
            n_feeds += 1
        if n_feeds:
            print(f"  ✓ {n_feeds} RSS feeds (*.xml)")

    # Selected Talks (top par qualité littéraire)
    selected_talks = collect_selected_talks(catalog)
    if selected_talks:
        selected_dir = SITE_DIR / "selected"
        selected_dir.mkdir(exist_ok=True)
        selected_page = build_selected_page(selected_talks)
        if selected_page:
            (selected_dir / "index.html").write_text(selected_page, encoding="utf-8")
        print(f"  ✓ selected/ ({len(selected_talks)} talks)")

    # Homepage
    (SITE_DIR / "index.html").write_text(build_homepage(catalog, selected_talks), encoding="utf-8")
    print(f"  ✓ index.html")

    # Pages de feeds et épisodes
    total_episodes = 0
    total_pdfs = 0
    prev_pdfs = 0
    for slug in FEED_ORDER:
        fdata = catalog.get(slug)
        if not fdata:
            continue
        feed_dir = SITE_DIR / slug
        feed_dir.mkdir()
        feed_name = fdata["name"]

        # Page feed
        (feed_dir / "index.html").write_text(
            build_feed_page(slug, fdata, catalog), encoding="utf-8"
        )

        # Collecter tous les épisodes valides (pour navigation prev/next)
        newest_first = slug in FEEDS_NEWEST_FIRST
        text_feed = is_text_feed(fdata)
        seasons_iter = list(fdata.get("seasons", []))
        if newest_first:
            seasons_iter = sorted(seasons_iter, key=_season_sort_key, reverse=True)
        valid_eps = []
        for season in seasons_iter:
            episodes = list(season.get("episodes", []))
            if newest_first:
                episodes = list(reversed(episodes))
            for ep in episodes:
                stem = ep_stem(ep)
                if not stem:
                    continue
                if text_feed:
                    if ep.get("word_count", 0) > 0:
                        valid_eps.append(ep)
                else:
                    if ep.get("duration_seconds", 0) > 0:
                        valid_eps.append(ep)

        # Pages épisodes
        for i, ep in enumerate(valid_eps):
            prev_ep = valid_eps[i - 1] if i > 0 else None
            next_ep = valid_eps[i + 1] if i < len(valid_eps) - 1 else None
            stem = ep_stem(ep)
            page_html = build_episode_page(slug, ep, prev_ep, next_ep, feed_name)
            (feed_dir / f"{stem}.html").write_text(page_html, encoding="utf-8")
            total_episodes += 1

            # PDF + EPUB + DOCX si un article beautifié existe
            article = load_article(slug, stem)
            if article:
                meta = load_metadata(slug, stem)
                title = meta.get("title_clean", ep.get("title", ""))
                if text_feed:
                    dur_str = format_reading_time(ep.get("reading_minutes", 0))
                else:
                    dur_str = format_duration(ep.get("duration_seconds", 0))
                clean_feed = feed_name.replace("Satipanya — ", "")
                desc = meta.get("description_long") or ep.get("description_long", "")
                desc = desc.replace(
                    "(This description was generated automatically, inaccuracies may happen in the process.)", ""
                ).strip()
                generate_pdf(
                    feed_dir / f"{stem}.pdf",
                    title, ep.get("speaker", ""), clean_feed, dur_str, article,
                )
                generate_epub(
                    feed_dir / f"{stem}.epub",
                    title, ep.get("speaker", ""), clean_feed, dur_str, article,
                    description=desc,
                )
                generate_docx(
                    feed_dir / f"{stem}.docx",
                    title, ep.get("speaker", ""), clean_feed, dur_str, article,
                )
                total_pdfs += 1

        print(f"  ✓ {slug}/ (index + {len(valid_eps)} episodes, {total_pdfs - prev_pdfs} PDFs/EPUBs/DOCXs)")
        prev_pdfs = total_pdfs

    # Search index
    search_data = build_search_index(catalog)
    (SITE_DIR / "search-index.json").write_text(
        json.dumps(search_data, ensure_ascii=False), encoding="utf-8"
    )
    index_size_kb = (SITE_DIR / "search-index.json").stat().st_size / 1024
    print(f"  ✓ search-index.json ({len(search_data)} entries, {index_size_kb:.0f} KB)")

    # Search page
    (SITE_DIR / "search.html").write_text(build_search_page(), encoding="utf-8")
    print(f"  ✓ search.html")

    # Topics page
    topics_index = build_topics_index(catalog)
    if topics_index:
        topics_dir = SITE_DIR / "topics"
        topics_dir.mkdir(exist_ok=True)
        (SITE_DIR / "topics-index.json").write_text(
            json.dumps(topics_index, ensure_ascii=False, separators=(",", ":")),
            encoding="utf-8",
        )
        topics_page = build_topics_page(topics_index)
        if topics_page:
            (topics_dir / "index.html").write_text(topics_page, encoding="utf-8")
        topics_size_kb = (SITE_DIR / "topics-index.json").stat().st_size / 1024
        print(f"  ✓ topics/ ({len(topics_index)} topics, index {topics_size_kb:.0f} KB)")

    print(f"\n{'=' * 60}")
    print(f"Site built: {total_episodes} episode pages, {total_pdfs} PDFs, 6 feed pages")
    print(f"Output: {SITE_DIR}/")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    build_site()
