#!/usr/bin/env python3
"""
build_books.py — Génère un livre PDF + EPUB pour chaque collection Satipanya.

Chaque feed produit un livre avec :
  - Page de couverture
  - Table des matières
  - Un chapitre par talk (résumé en accroche + transcript complet)

Utilise WeasyPrint (CSS paged media) pour le PDF A4 et ebooklib pour l'EPUB.
Doit être exécuté avec le conda env 'newspapers' :
    conda run -n newspapers python build_books.py

Résultat dans site/books/
"""

import json
import re
from pathlib import Path
from html import escape as esc

from weasyprint import HTML
from ebooklib import epub

# ── Configuration ──────────────────────────────────────────────

PROJECT_DIR = Path(__file__).parent
CATALOG_PATH = PROJECT_DIR / "catalog.json"
METADATA_DIR = PROJECT_DIR / "metadata"
ARTICLES_DIR = PROJECT_DIR / "articles"
BOOKS_DIR    = PROJECT_DIR / "site" / "books"

FEED_ORDER = [
    "guided-meditations",
    "foundation-course",
    "dhammabytes",
    "dharma-talks",
    "noirins-teachings",
    "international-talks",
    "youtube-talks",
    "bhante-essays",
    "noirin-essays",
    "tips-of-the-day",
]

FEED_SUBTITLE = {
    "guided-meditations": "Guided vipassanā meditation practices and Pāli chanting",
    "foundation-course": "A systematic introduction to Theravāda Buddhist thought and practice",
    "dhammabytes": "Short teachings on key Buddhist concepts and doctrines",
    "dharma-talks": "Retreat talks on meditation, ethics, and the path to liberation",
    "noirins-teachings": "Dharma teachings by Nóirín Ní Riain",
    "international-talks": "Talks by visiting teachers from the Theravāda tradition",
    "youtube-talks": "Dharma talks and teachings from the Satipanya YouTube channel",
    "bhante-essays": "Written essays and teachings on Buddhist philosophy and practice",
    "noirin-essays": "Written essays and reflections on practice and ethics",
    "tips-of-the-day": "Short practical tips for integrating mindfulness into daily life",
}


# ── Utilitaires ────────────────────────────────────────────────

def format_duration(seconds):
    if not seconds:
        return ""
    s = int(seconds)
    h, s = divmod(s, 3600)
    m, s = divmod(s, 60)
    return f"{h}h{m:02d}" if h else f"{m} min"


def load_metadata(feed_slug, stem):
    path = METADATA_DIR / feed_slug / f"{stem}.json"
    if path.exists():
        with open(path) as f:
            return json.load(f)
    return {}


def load_article(feed_slug, stem):
    path = ARTICLES_DIR / feed_slug / f"{stem}.txt"
    if path.exists():
        return path.read_text(encoding="utf-8").strip()
    return None


def ep_stem(ep):
    if ep.get("stem"):
        return ep["stem"]
    tp = ep.get("transcript_path")
    return Path(tp).stem if tp else None


def is_text_episode(ep):
    return ep.get("content_type") == "text"


def is_text_feed(fdata):
    return fdata.get("content_type") == "text"


def format_reading_time(minutes):
    if not minutes:
        return ""
    return f"{int(minutes)} min read"


def article_to_html(text):
    """Convertit un article texte brut en HTML avec paragraphes et italiques."""
    paragraphs = text.split("\n\n")
    parts = []
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        p = esc(p)
        # *italiques* → <em>
        p = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', p)
        p = p.replace("\n", "<br>")
        parts.append(f"<p>{p}</p>")
    return "\n".join(parts)


def clean_description(text):
    """Retire la note auto-générée de la description."""
    return text.replace(
        "(This description was generated automatically, inaccuracies may happen in the process.)", ""
    ).strip()


# ── CSS pour le PDF A4 ────────────────────────────────────────

PDF_CSS = """\
@page {
    size: A4;
    margin: 28mm 25mm 30mm 30mm;
    @bottom-center {
        content: counter(page);
        font-family: "Noto Sans", sans-serif;
        font-size: 8pt;
        color: #999;
    }
    @top-center {
        content: string(book-title);
        font-family: "Noto Serif", serif;
        font-size: 7.5pt;
        font-style: italic;
        color: #aaa;
        letter-spacing: 0.03em;
    }
}
@page :first {
    margin: 0;
    @bottom-center { content: none; }
    @top-center { content: none; }
}
@page :blank {
    @bottom-center { content: none; }
    @top-center { content: none; }
}
@page chapter-first {
    @top-center { content: none; }
}

* { margin: 0; padding: 0; box-sizing: border-box; }

body {
    font-family: "Noto Serif", Georgia, serif;
    font-size: 10.5pt;
    line-height: 1.65;
    color: #1a1a1a;
    text-align: justify;
    hyphens: auto;
    -webkit-hyphens: auto;
    orphans: 3;
    widows: 3;
}

/* ── Couverture ─────────────────────────────────── */
.cover {
    page-break-after: always;
    width: 210mm; height: 297mm;
    display: flex; flex-direction: column;
    justify-content: center; align-items: center;
    text-align: center;
    background: #faf9f7;
    padding: 40mm 30mm;
    position: relative;
}
.cover::before {
    content: ""; position: absolute;
    top: 25mm; left: 30mm; right: 30mm;
    height: 0.6pt; background: #B8860B;
}
.cover::after {
    content: ""; position: absolute;
    bottom: 25mm; left: 30mm; right: 30mm;
    height: 0.6pt; background: #B8860B;
}
.cover-series {
    font-family: "Noto Sans", sans-serif;
    font-size: 9pt;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.2em;
    color: #B8860B;
    margin-bottom: 8mm;
}
.cover-title {
    font-family: "Noto Serif", serif;
    font-size: 28pt;
    font-weight: 700;
    line-height: 1.2;
    color: #1a1a1a;
    margin-bottom: 6mm;
    letter-spacing: -0.01em;
    hyphens: none;
}
.cover-subtitle {
    font-family: "Noto Serif", serif;
    font-size: 11pt;
    font-style: italic;
    color: #666;
    max-width: 120mm;
    line-height: 1.6;
    margin-bottom: 15mm;
}
.cover-author {
    font-family: "Noto Sans", sans-serif;
    font-size: 11pt;
    font-weight: 600;
    color: #333;
    letter-spacing: 0.05em;
    margin-bottom: 4mm;
}
.cover-retreat {
    font-family: "Noto Sans", sans-serif;
    font-size: 8.5pt;
    color: #999;
    letter-spacing: 0.08em;
    text-transform: uppercase;
}

/* ── Titre courant ──────────────────────────────── */
h1.book-title-string {
    string-set: book-title content();
    font-size: 0; height: 0; margin: 0; padding: 0;
    visibility: hidden;
}

/* ── Table des matières ─────────────────────────── */
.toc-page {
    page-break-after: always;
    padding-top: 15mm;
}
.toc-page h2 {
    font-family: "Noto Serif", serif;
    font-size: 18pt;
    font-weight: 700;
    margin-bottom: 10mm;
    color: #1a1a1a;
    letter-spacing: -0.01em;
}
.toc-season {
    font-family: "Noto Sans", sans-serif;
    font-size: 8pt;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.15em;
    color: #B8860B;
    margin-top: 6mm;
    margin-bottom: 2mm;
    padding-bottom: 1.5mm;
    border-bottom: 0.4pt solid #e0d8cf;
}
.toc-entry {
    display: flex;
    justify-content: space-between;
    align-items: baseline;
    padding: 1.2mm 0;
    border-bottom: 0.2pt dotted #ddd;
    text-decoration: none;
    color: #1a1a1a;
}
.toc-entry:last-child { border-bottom: none; }
.toc-title {
    font-family: "Noto Serif", serif;
    font-size: 9.5pt;
    flex: 1;
    padding-right: 3mm;
}
.toc-duration {
    font-family: "Noto Sans", sans-serif;
    font-size: 7.5pt;
    color: #999;
    white-space: nowrap;
    padding-right: 3mm;
}
.toc-entry::after {
    content: target-counter(attr(href url), page);
    font-family: "Noto Sans", sans-serif;
    font-size: 8pt;
    color: #999;
    white-space: nowrap;
    min-width: 8mm;
    text-align: right;
}

/* ── Chapitres ──────────────────────────────────── */
.chapter {
    page-break-before: always;
    page: chapter-first;
}
.chapter-header {
    margin-bottom: 8mm;
    padding-bottom: 5mm;
    border-bottom: 0.5pt solid #B8860B;
}
.chapter-number {
    font-family: "Noto Sans", sans-serif;
    font-size: 8pt;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.15em;
    color: #B8860B;
    margin-bottom: 2mm;
}
.chapter-title {
    font-family: "Noto Serif", serif;
    font-size: 18pt;
    font-weight: 700;
    line-height: 1.25;
    color: #1a1a1a;
    margin-bottom: 2mm;
    letter-spacing: -0.01em;
}
.chapter-meta {
    font-family: "Noto Sans", sans-serif;
    font-size: 8pt;
    color: #999;
}

/* Accroche (résumé) */
.chapter-lead {
    font-family: "Noto Serif", serif;
    font-size: 10pt;
    font-style: italic;
    color: #555;
    line-height: 1.7;
    margin-bottom: 6mm;
    padding-left: 4mm;
    border-left: 2pt solid #e0d8cf;
}
.chapter-lead p { margin-bottom: 3mm; }
.chapter-lead p:last-child { margin-bottom: 0; }

/* Corps du transcript */
.chapter-body {
    font-size: 10.5pt;
    line-height: 1.65;
}
.chapter-body p {
    margin-bottom: 3.5mm;
    text-indent: 0;
}
.chapter-body p + p {
    text-indent: 5mm;
}
.chapter-body p:first-child {
    text-indent: 0;
}
.chapter-body em {
    font-style: italic;
}

/* Premier paragraphe : pas d'indentation */
.chapter-body .first-para {
    text-indent: 0;
}

/* ── Colophon ───────────────────────────────────── */
.colophon {
    page-break-before: always;
    padding-top: 60mm;
    text-align: center;
}
.colophon p {
    font-family: "Noto Sans", sans-serif;
    font-size: 8pt;
    color: #999;
    line-height: 1.8;
}
.colophon .retreat-name {
    font-family: "Noto Serif", serif;
    font-size: 12pt;
    color: #333;
    margin-bottom: 3mm;
    font-weight: 600;
}
.colophon .colophon-rule {
    width: 30mm; height: 0.4pt;
    background: #B8860B;
    margin: 8mm auto;
}
"""


# ── Génération PDF ─────────────────────────────────────────────

def build_pdf_book(slug, fdata, output_path):
    """Génère un livre PDF A4 pour une collection."""
    name = fdata["name"].replace("Satipanya — ", "")
    subtitle = FEED_SUBTITLE.get(slug, "")

    # Collecter les auteurs
    authors = set()
    for season in fdata.get("seasons", []):
        for ep in season.get("episodes", []):
            if ep.get("speaker"):
                authors.add(ep["speaker"])
    author_str = ", ".join(sorted(authors))

    # ── Couverture ──
    cover_html = f"""
    <div class="cover">
        <div class="cover-series">Satipanya Buddhist Retreat</div>
        <div class="cover-title">{esc(name)}</div>
        <div class="cover-subtitle">{esc(subtitle)}</div>
        <div class="cover-author">{esc(author_str)}</div>
        <div class="cover-retreat">Shropshire, Wales · United Kingdom</div>
    </div>
    <h1 class="book-title-string">{esc(name)}</h1>
    """

    # ── Collecter les chapitres ──
    text_feed = is_text_feed(fdata)
    chapters = []
    for season in fdata.get("seasons", []):
        for ep in season.get("episodes", []):
            stem = ep_stem(ep)
            if not stem:
                continue
            if text_feed:
                if ep.get("word_count", 0) == 0:
                    continue
                dur_str = format_reading_time(ep.get("reading_minutes", 0))
            else:
                dur = ep.get("duration_seconds", 0)
                if dur == 0:
                    continue
                dur_str = format_duration(dur)
            article = load_article(slug, stem)
            if not article:
                continue  # pas de transcript beautifié → pas de chapitre
            meta = load_metadata(slug, stem)
            title = meta.get("title_clean", ep.get("title", "Untitled"))
            desc = clean_description(
                meta.get("description_long") or ep.get("description_long", "")
            )
            chapters.append({
                "title": title,
                "speaker": ep.get("speaker", ""),
                "duration": dur_str,
                "description": desc,
                "article_html": article_to_html(article),
                "season_name": season.get("name", ""),
                "season_number": season.get("number", 1),
            })

    if not chapters:
        return 0

    # ── Table des matières ──
    toc_html = '<div class="toc-page"><h2>Contents</h2>\n'
    current_season = None
    for i, ch in enumerate(chapters):
        if ch["season_name"] != current_season:
            current_season = ch["season_name"]
            toc_html += f'<div class="toc-season">{esc(current_season)}</div>\n'
        toc_html += f"""
        <a class="toc-entry" href="#ch-{i}">
            <span class="toc-title">{esc(ch["title"])}</span>
            <span class="toc-duration">{ch["duration"]}</span>
        </a>\n"""
    toc_html += "</div>\n"

    # ── Chapitres ──
    chapters_html = ""
    for i, ch in enumerate(chapters):
        # Ajouter la classe first-para au premier <p> du body
        body = ch["article_html"]
        body = body.replace("<p>", '<p class="first-para">', 1)

        # Accroche (description)
        lead_html = ""
        if ch["description"]:
            desc_paras = [f"<p>{esc(p.strip())}</p>"
                          for p in ch["description"].split("\n\n") if p.strip()]
            lead_html = f'<div class="chapter-lead">{"".join(desc_paras)}</div>'

        chapters_html += f"""
        <section class="chapter" id="ch-{i}">
            <div class="chapter-header">
                <div class="chapter-number">Chapter {i + 1}</div>
                <div class="chapter-title">{esc(ch["title"])}</div>
                <div class="chapter-meta">{esc(ch["speaker"])} · {ch["duration"]}</div>
            </div>
            {lead_html}
            <div class="chapter-body">
                {body}
            </div>
        </section>
        """

    # ── Colophon ──
    content_word = "essays" if text_feed else "talks"
    provenance = ("Published on satipanya.org.uk." if text_feed else
                  "Transcriptions produced locally using Swiss low-carbon electricity.<br>"
                  "Corrections and rewriting by cloud-hosted AI.")
    colophon_html = f"""
    <div class="colophon">
        <div class="retreat-name">Satipanya Buddhist Retreat</div>
        <p>{esc(name)}<br>{esc(subtitle)}</p>
        <div class="colophon-rule"></div>
        <p>{len(chapters)} {content_word} · {esc(author_str)}</p>
        <p style="margin-top: 5mm;">
            {provenance}
        </p>
        <p style="margin-top: 5mm;">
            <a href="https://www.satipanya.org.uk">satipanya.org.uk</a>
        </p>
    </div>
    """

    # ── Assemblage HTML final ──
    full_html = f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><style>{PDF_CSS}</style></head>
<body>
{cover_html}
{toc_html}
{chapters_html}
{colophon_html}
</body></html>"""

    HTML(string=full_html).write_pdf(str(output_path))
    return len(chapters)


# ── Génération EPUB ────────────────────────────────────────────

EPUB_CSS = """\
body {
    font-family: Georgia, "Times New Roman", serif;
    font-size: 1em;
    line-height: 1.7;
    color: #1a1a1a;
    margin: 0;
    padding: 0;
}
h1 {
    font-size: 1.6em;
    font-weight: 700;
    line-height: 1.25;
    margin: 0 0 0.3em;
    color: #1a1a1a;
}
.chapter-meta {
    font-size: 0.8em;
    color: #999;
    margin-bottom: 1em;
}
.chapter-lead {
    font-style: italic;
    color: #555;
    border-left: 3px solid #e0d8cf;
    padding-left: 1em;
    margin-bottom: 1.5em;
    line-height: 1.7;
}
.chapter-lead p { margin-bottom: 0.5em; }
.chapter-body p {
    margin-bottom: 0.8em;
    text-align: justify;
}
.chapter-body em { font-style: italic; }
.colophon {
    text-align: center;
    margin-top: 3em;
    color: #999;
    font-size: 0.85em;
    line-height: 1.8;
}
"""


def build_epub_book(slug, fdata, output_path):
    """Génère un livre EPUB pour une collection."""
    name = fdata["name"].replace("Satipanya — ", "")
    subtitle = FEED_SUBTITLE.get(slug, "")

    authors = set()
    for season in fdata.get("seasons", []):
        for ep in season.get("episodes", []):
            if ep.get("speaker"):
                authors.add(ep["speaker"])
    author_str = ", ".join(sorted(authors))

    book = epub.EpubBook()
    book.set_identifier(f"satipanya-{slug}")
    book.set_title(f"{name} — Satipanya Buddhist Retreat")
    book.set_language("en")
    for author in sorted(authors):
        book.add_author(author)
    book.add_metadata("DC", "publisher", "Satipanya Buddhist Retreat")
    book.add_metadata("DC", "description", subtitle)

    # CSS
    style = epub.EpubItem(uid="style", file_name="style/default.css",
                          media_type="text/css", content=EPUB_CSS.encode())
    book.add_item(style)

    # Collecter chapitres
    text_feed = is_text_feed(fdata)
    epub_chapters = []
    toc_entries = []
    spine = ["nav"]

    chapter_idx = 0
    for season in fdata.get("seasons", []):
        season_chapters = []
        for ep in season.get("episodes", []):
            stem = ep_stem(ep)
            if not stem:
                continue
            if text_feed:
                if ep.get("word_count", 0) == 0:
                    continue
                dur_str = format_reading_time(ep.get("reading_minutes", 0))
            else:
                dur = ep.get("duration_seconds", 0)
                if dur == 0:
                    continue
                dur_str = format_duration(dur)
            article = load_article(slug, stem)
            if not article:
                continue
            meta = load_metadata(slug, stem)
            title = meta.get("title_clean", ep.get("title", "Untitled"))
            desc = clean_description(
                meta.get("description_long") or ep.get("description_long", "")
            )
            speaker = ep.get("speaker", "")

            # Accroche
            lead_html = ""
            if desc:
                desc_paras = "".join(
                    f"<p>{esc(p.strip())}</p>"
                    for p in desc.split("\n\n") if p.strip()
                )
                lead_html = f'<div class="chapter-lead">{desc_paras}</div>'

            body_html = article_to_html(article)

            ch = epub.EpubHtml(
                title=title,
                file_name=f"ch{chapter_idx:03d}.xhtml",
                lang="en",
            )
            ch.content = f"""<html><head></head><body>
<h1>{esc(title)}</h1>
<div class="chapter-meta">{esc(speaker)} · {dur_str}</div>
{lead_html}
<div class="chapter-body">{body_html}</div>
</body></html>"""
            ch.add_item(style)
            book.add_item(ch)
            epub_chapters.append(ch)
            season_chapters.append(ch)
            spine.append(ch)
            chapter_idx += 1

        if season_chapters:
            section = epub.Section(season.get("name", f"Season {season.get('number', '?')}"))
            toc_entries.append((section, season_chapters))

    if not epub_chapters:
        return 0

    # Colophon
    content_word = "essays" if text_feed else "talks"
    provenance = ("Published on satipanya.org.uk." if text_feed else
                  "Transcriptions produced locally using Swiss low-carbon electricity. "
                  "Corrections and rewriting by cloud-hosted AI.")
    colophon = epub.EpubHtml(title="About", file_name="colophon.xhtml", lang="en")
    colophon.content = f"""<html><head></head><body>
<div class="colophon">
<p><strong>Satipanya Buddhist Retreat</strong></p>
<p>{esc(name)}</p>
<p>{esc(subtitle)}</p>
<p>{len(epub_chapters)} {content_word} · {esc(author_str)}</p>
<p>{provenance}</p>
<p><a href="https://www.satipanya.org.uk">satipanya.org.uk</a></p>
</div></body></html>"""
    colophon.add_item(style)
    book.add_item(colophon)
    spine.append(colophon)

    book.toc = toc_entries
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = spine

    epub.write_epub(str(output_path), book, {})
    return len(epub_chapters)


# ── Génération DOCX ───────────────────────────────────────────

def _add_article_paragraphs(doc, text):
    """Ajoute les paragraphes d'un article au document DOCX avec support *italique*."""
    paragraphs = text.split("\n\n")
    for p_text in paragraphs:
        p_text = p_text.strip()
        if not p_text:
            continue
        p = doc.add_paragraph()
        parts = re.split(r'(\*[^*]+\*)', p_text)
        for part in parts:
            if part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                p.add_run(part.replace("\n", " "))


def build_docx_book(slug, fdata, output_path):
    """Génère un livre DOCX pour une collection."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    name = fdata["name"].replace("Satipanya — ", "")
    subtitle = FEED_SUBTITLE.get(slug, "")

    authors = set()
    for season in fdata.get("seasons", []):
        for ep in season.get("episodes", []):
            if ep.get("speaker"):
                authors.add(ep["speaker"])
    author_str = ", ".join(sorted(authors))

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.8)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ── Couverture ──
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SATIPANYA BUDDHIST RETREAT")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
    run.font.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.size = Pt(28)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(author_str)
    run.font.size = Pt(11)
    run.font.bold = True

    doc.add_page_break()

    # ── Collecter les chapitres ──
    text_feed = is_text_feed(fdata)
    chapters = []
    for season in fdata.get("seasons", []):
        for ep in season.get("episodes", []):
            stem = ep_stem(ep)
            if not stem:
                continue
            if text_feed:
                if ep.get("word_count", 0) == 0:
                    continue
                dur_str = format_reading_time(ep.get("reading_minutes", 0))
            else:
                dur = ep.get("duration_seconds", 0)
                if dur == 0:
                    continue
                dur_str = format_duration(dur)
            article = load_article(slug, stem)
            if not article:
                continue
            meta = load_metadata(slug, stem)
            title = meta.get("title_clean", ep.get("title", "Untitled"))
            desc = clean_description(
                meta.get("description_long") or ep.get("description_long", "")
            )
            chapters.append({
                "title": title,
                "speaker": ep.get("speaker", ""),
                "duration": dur_str,
                "description": desc,
                "article": article,
                "season_name": season.get("name", ""),
            })

    if not chapters:
        return 0

    # ── Table des matières ──
    p = doc.add_paragraph()
    run = p.add_run("Contents")
    run.font.size = Pt(18)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(12)

    current_season = None
    for ch in chapters:
        if ch["season_name"] != current_season:
            current_season = ch["season_name"]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(current_season.upper())
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(ch["title"])
        run.font.size = Pt(9.5)
        if ch["duration"]:
            run = p.add_run(f"   {ch['duration']}")
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ── Chapitres ──
    for i, ch in enumerate(chapters):
        if i > 0:
            doc.add_page_break()

        p = doc.add_paragraph()
        run = p.add_run(f"CHAPTER {i + 1}")
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

        p = doc.add_paragraph()
        run = p.add_run(ch["title"])
        run.font.size = Pt(18)
        run.font.bold = True

        p = doc.add_paragraph()
        run = p.add_run(f"{ch['speaker']} · {ch['duration']}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        p.paragraph_format.space_after = Pt(10)

        if ch["description"]:
            for desc_para in ch["description"].split("\n\n"):
                desc_para = desc_para.strip()
                if desc_para:
                    p = doc.add_paragraph()
                    run = p.add_run(desc_para)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        _add_article_paragraphs(doc, ch["article"])

    # ── Colophon ──
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Satipanya Buddhist Retreat")
    run.font.size = Pt(12)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{name}\n{subtitle}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    content_word = "essays" if text_feed else "talks"
    provenance = ("Published on satipanya.org.uk." if text_feed else
                  "Transcriptions produced locally using Swiss low-carbon electricity.\n"
                  "Corrections and rewriting by cloud-hosted AI.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{len(chapters)} {content_word} · {author_str}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(provenance)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(str(output_path))
    return len(chapters)


# ── Selected Talks book ────────────────────────────────────────

SELECTED_TALKS_COUNT = 120

def collect_selected_chapters(catalog):
    """Collecte les chapitres pour le livre Selected Talks.

    Alterne les enseignants dans les premières positions.
    """
    scored = []
    for slug in FEED_ORDER:
        fdata = catalog.get(slug)
        if not fdata:
            continue
        feed_name = fdata["name"].replace("Satipanya — ", "")
        text_feed = is_text_feed(fdata)
        for season in fdata.get("seasons", []):
            for ep in season.get("episodes", []):
                score = ep.get("lite_score")
                if score is None or score < 1:
                    continue
                stem = ep_stem(ep)
                if not stem:
                    continue
                if text_feed:
                    if ep.get("word_count", 0) == 0:
                        continue
                    dur_str = format_reading_time(ep.get("reading_minutes", 0))
                else:
                    if ep.get("duration_seconds", 0) == 0:
                        continue
                    dur_str = format_duration(ep.get("duration_seconds", 0))
                article = load_article(slug, stem)
                if not article:
                    continue
                meta = load_metadata(slug, stem)
                title = meta.get("title_clean", ep.get("title", "Untitled"))
                desc = clean_description(
                    meta.get("description_long") or ep.get("description_long", "")
                )
                scored.append({
                    "title": title,
                    "speaker": ep.get("speaker", ""),
                    "duration": dur_str,
                    "description": desc,
                    "article": article,
                    "article_html": article_to_html(article),
                    "feed_name": feed_name,
                    "score": score,
                    "season_name": feed_name,
                    "season_number": 0,
                })

    scored.sort(key=lambda x: (-x["score"], x["title"]))

    # Alterner les enseignants dans les 10 premières positions
    ALTERNATE_COUNT = 10
    if len(scored) > ALTERNATE_COUNT:
        noirin = [t for t in scored if "noirin" in t["speaker"].lower()]
        bhante = [t for t in scored if "noirin" not in t["speaker"].lower()]
        alternated = []
        ni, bi = 0, 0
        for pos in range(ALTERNATE_COUNT):
            if pos % 2 == 0:
                if bi < len(bhante):
                    alternated.append(bhante[bi])
                    bi += 1
                elif ni < len(noirin):
                    alternated.append(noirin[ni])
                    ni += 1
            else:
                if ni < len(noirin):
                    alternated.append(noirin[ni])
                    ni += 1
                elif bi < len(bhante):
                    alternated.append(bhante[bi])
                    bi += 1
        used = set(id(t) for t in alternated)
        rest = [t for t in scored if id(t) not in used]
        scored = alternated + rest

    return scored[:SELECTED_TALKS_COUNT]


def build_selected_pdf(chapters, output_path):
    """Génère le PDF pour le livre Selected Talks."""
    name = "Selected Talks"
    subtitle = "A selection of interesting talks if you don't know where to start"
    authors = set(ch["speaker"] for ch in chapters if ch["speaker"])
    author_str = ", ".join(sorted(authors))

    cover_html = f"""
    <div class="cover">
        <div class="cover-series">Satipanya Buddhist Retreat</div>
        <div class="cover-title">{esc(name)}</div>
        <div class="cover-subtitle">{esc(subtitle)}</div>
        <div class="cover-author">{esc(author_str)}</div>
        <div class="cover-retreat">Shropshire, Wales · United Kingdom</div>
    </div>
    <h1 class="book-title-string">{esc(name)}</h1>
    """

    toc_html = '<div class="toc-page"><h2>Contents</h2>\n'
    for i, ch in enumerate(chapters):
        toc_html += f"""
        <a class="toc-entry" href="#ch-{i}">
            <span class="toc-title">{esc(ch["title"])}</span>
            <span class="toc-duration">{ch["duration"]}</span>
        </a>\n"""
    toc_html += "</div>\n"

    chapters_html = ""
    for i, ch in enumerate(chapters):
        body = ch["article_html"]
        body = body.replace("<p>", '<p class="first-para">', 1)
        lead_html = ""
        if ch["description"]:
            desc_paras = [f"<p>{esc(p.strip())}</p>"
                          for p in ch["description"].split("\n\n") if p.strip()]
            lead_html = f'<div class="chapter-lead">{"".join(desc_paras)}</div>'
        chapters_html += f"""
        <section class="chapter" id="ch-{i}">
            <div class="chapter-header">
                <div class="chapter-number">Chapter {i + 1}</div>
                <div class="chapter-title">{esc(ch["title"])}</div>
                <div class="chapter-meta">{esc(ch["speaker"])} · {ch["duration"]}</div>
            </div>
            {lead_html}
            <div class="chapter-body">{body}</div>
        </section>
        """

    colophon_html = f"""
    <div class="colophon">
        <div class="retreat-name">Satipanya Buddhist Retreat</div>
        <p>{esc(name)}<br>{esc(subtitle)}</p>
        <div class="colophon-rule"></div>
        <p>{len(chapters)} talks · {esc(author_str)}</p>
        <p style="margin-top: 5mm;">
            Transcriptions produced locally using Swiss low-carbon electricity.<br>
            Corrections and rewriting by cloud-hosted AI.
        </p>
        <p style="margin-top: 5mm;">
            <a href="https://www.satipanya.org.uk">satipanya.org.uk</a>
        </p>
    </div>
    """

    full_html = f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><style>{PDF_CSS}</style></head>
<body>
{cover_html}
{toc_html}
{chapters_html}
{colophon_html}
</body></html>"""

    HTML(string=full_html).write_pdf(str(output_path))
    return len(chapters)


def build_selected_epub(chapters, output_path):
    """Génère l'EPUB pour le livre Selected Talks."""
    name = "Selected Talks"
    subtitle = "A selection of interesting talks if you don't know where to start"
    authors = set(ch["speaker"] for ch in chapters if ch["speaker"])
    author_str = ", ".join(sorted(authors))

    book = epub.EpubBook()
    book.set_identifier("satipanya-selected-talks")
    book.set_title(f"{name} — Satipanya Buddhist Retreat")
    book.set_language("en")
    for author in sorted(authors):
        book.add_author(author)
    book.add_metadata("DC", "publisher", "Satipanya Buddhist Retreat")
    book.add_metadata("DC", "description", subtitle)

    style = epub.EpubItem(uid="style", file_name="style/default.css",
                          media_type="text/css", content=EPUB_CSS.encode())
    book.add_item(style)

    epub_chapters = []
    spine = ["nav"]

    for i, ch_data in enumerate(chapters):
        title = ch_data["title"]
        speaker = ch_data["speaker"]
        dur_str = ch_data["duration"]
        desc = ch_data["description"]
        body_html = ch_data["article_html"]

        lead_html = ""
        if desc:
            desc_paras = "".join(
                f"<p>{esc(p.strip())}</p>"
                for p in desc.split("\n\n") if p.strip()
            )
            lead_html = f'<div class="chapter-lead">{desc_paras}</div>'

        ch = epub.EpubHtml(title=title, file_name=f"ch{i:03d}.xhtml", lang="en")
        ch.content = f"""<html><head></head><body>
<h1>{esc(title)}</h1>
<div class="chapter-meta">{esc(speaker)} · {dur_str}</div>
{lead_html}
<div class="chapter-body">{body_html}</div>
</body></html>"""
        ch.add_item(style)
        book.add_item(ch)
        epub_chapters.append(ch)
        spine.append(ch)

    if not epub_chapters:
        return 0

    colophon = epub.EpubHtml(title="About", file_name="colophon.xhtml", lang="en")
    colophon.content = f"""<html><head></head><body>
<div class="colophon">
<p><strong>Satipanya Buddhist Retreat</strong></p>
<p>{esc(name)}</p>
<p>{esc(subtitle)}</p>
<p>{len(epub_chapters)} talks · {esc(author_str)}</p>
<p>Transcriptions produced locally using Swiss low-carbon electricity.
Corrections and rewriting by cloud-hosted AI.</p>
<p><a href="https://www.satipanya.org.uk">satipanya.org.uk</a></p>
</div></body></html>"""
    colophon.add_item(style)
    book.add_item(colophon)
    spine.append(colophon)

    book.toc = epub_chapters
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = spine

    epub.write_epub(str(output_path), book, {})
    return len(epub_chapters)


def build_selected_docx(chapters, output_path):
    """Génère le DOCX pour le livre Selected Talks."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    name = "Selected Talks"
    subtitle = "A selection of interesting talks if you don't know where to start"
    authors = set(ch["speaker"] for ch in chapters if ch["speaker"])
    author_str = ", ".join(sorted(authors))

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.8)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SATIPANYA BUDDHIST RETREAT")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
    run.font.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.size = Pt(28)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(author_str)
    run.font.size = Pt(11)
    run.font.bold = True

    doc.add_page_break()

    # Table des matières
    p = doc.add_paragraph()
    run = p.add_run("Contents")
    run.font.size = Pt(18)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(12)

    for ch in chapters:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(ch["title"])
        run.font.size = Pt(9.5)
        if ch["duration"]:
            run = p.add_run(f"   {ch['duration']}")
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # Chapitres
    for i, ch in enumerate(chapters):
        if i > 0:
            doc.add_page_break()

        p = doc.add_paragraph()
        run = p.add_run(f"CHAPTER {i + 1}")
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

        p = doc.add_paragraph()
        run = p.add_run(ch["title"])
        run.font.size = Pt(18)
        run.font.bold = True

        p = doc.add_paragraph()
        run = p.add_run(f"{ch['speaker']} · {ch['duration']}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        p.paragraph_format.space_after = Pt(10)

        if ch["description"]:
            for desc_para in ch["description"].split("\n\n"):
                desc_para = desc_para.strip()
                if desc_para:
                    p = doc.add_paragraph()
                    run = p.add_run(desc_para)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        _add_article_paragraphs(doc, ch["article"])

    # Colophon
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Satipanya Buddhist Retreat")
    run.font.size = Pt(12)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{name}\n{subtitle}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{len(chapters)} talks · {author_str}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Transcriptions produced locally using Swiss low-carbon electricity.\n"
        "Corrections and rewriting by cloud-hosted AI."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(str(output_path))
    return len(chapters)


# ── Main ───────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("Building Satipanya book collection (PDF + EPUB + DOCX)")
    print("=" * 60)

    with open(CATALOG_PATH) as f:
        raw_catalog = json.load(f)

    catalog = {}
    for key, fdata in raw_catalog.items():
        slug = fdata.get("slug", key.replace("_", "-"))
        catalog[slug] = fdata

    BOOKS_DIR.mkdir(parents=True, exist_ok=True)

    for slug in FEED_ORDER:
        fdata = catalog.get(slug)
        if not fdata:
            continue
        name = fdata["name"].replace("Satipanya — ", "")

        # PDF
        pdf_path = BOOKS_DIR / f"{slug}.pdf"
        print(f"\n  {name}...")
        n_pdf = build_pdf_book(slug, fdata, pdf_path)
        if n_pdf:
            size_mb = pdf_path.stat().st_size / (1024 * 1024)
            print(f"    ✓ PDF: {n_pdf} chapters, {size_mb:.1f} MB")
        else:
            print(f"    ⏭ PDF: no beautified transcripts yet")

        # EPUB
        epub_path = BOOKS_DIR / f"{slug}.epub"
        n_epub = build_epub_book(slug, fdata, epub_path)
        if n_epub:
            size_mb = epub_path.stat().st_size / (1024 * 1024)
            print(f"    ✓ EPUB: {n_epub} chapters, {size_mb:.1f} MB")
        else:
            if epub_path.exists():
                epub_path.unlink()
            print(f"    ⏭ EPUB: no beautified transcripts yet")

        # DOCX
        docx_path = BOOKS_DIR / f"{slug}.docx"
        n_docx = build_docx_book(slug, fdata, docx_path)
        if n_docx:
            size_mb = docx_path.stat().st_size / (1024 * 1024)
            print(f"    ✓ DOCX: {n_docx} chapters, {size_mb:.1f} MB")
        else:
            if docx_path.exists():
                docx_path.unlink()
            print(f"    ⏭ DOCX: no beautified transcripts yet")

    # ── Selected Talks book ──
    print(f"\n  Selected Talks...")
    selected_chapters = collect_selected_chapters(catalog)
    if selected_chapters:
        pdf_path = BOOKS_DIR / "selected-talks.pdf"
        n = build_selected_pdf(selected_chapters, pdf_path)
        size_mb = pdf_path.stat().st_size / (1024 * 1024)
        print(f"    ✓ PDF: {n} chapters, {size_mb:.1f} MB")

        epub_path = BOOKS_DIR / "selected-talks.epub"
        n = build_selected_epub(selected_chapters, epub_path)
        size_mb = epub_path.stat().st_size / (1024 * 1024)
        print(f"    ✓ EPUB: {n} chapters, {size_mb:.1f} MB")

        docx_path = BOOKS_DIR / "selected-talks.docx"
        n = build_selected_docx(selected_chapters, docx_path)
        size_mb = docx_path.stat().st_size / (1024 * 1024)
        print(f"    ✓ DOCX: {n} chapters, {size_mb:.1f} MB")
    else:
        print(f"    ⏭ no scored talks with beautified transcripts")

    print(f"\n{'=' * 60}")
    print(f"Books output: {BOOKS_DIR}/")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
