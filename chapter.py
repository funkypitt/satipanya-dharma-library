#!/usr/bin/env python3
"""
chapter.py — Compose un chapitre de livre à partir des enseignements de
Bhante Bodhidhamma sur un thème donné.

Le pipeline :
  1. relevance — Score chaque épisode de Bhante (audio + texte) sur
     sa pertinence par rapport au thème (métadonnées uniquement, Sonnet).
  2. select    — Claude Opus choisit LE discours de base le plus complet
     et ayant le plus de potentiel littéraire parmi les top candidats.
  4. compose   — Opus structure le discours de base en chapitre JSON
     (titre, sous-titre, épigraphe, segments ordonnés).
  4c. prune    — Opus supprime les redites, boucles orales, digressions.
  5. render    — Génération d'un .docx (texte noir uniforme).
  5b. preprint — DOCX prêt à imprimer (typo soignée, logo).
  6. refine    — Polissage littéraire léger (oral → écrit).
  6b. render   — DOCX preprint de la version polie.
  7. heavier   — Polissage éditorial plus profond (qualité livre).
  7b. render   — DOCX preprint de la version livre.

Checkpoints sauvegardés dans chapters/<slug>/ :
  - candidates.json, base.json, composition.json, pruned.json,
    refined.json, refined_heavier.json, chapter.docx,
    chapter-preprint.docx, chapter-preprint-refined.docx,
    chapter-preprint-refined-heavier.docx

Usage :
    python chapter.py --theme "Dharma in Daily Life"
    python chapter.py --theme "..." --top-candidates 12

Variables d'environnement requises :
    ANTHROPIC_API_KEY

Environnement conda recommandé : interview
"""

import argparse
import json
import os
import re
import sys
import time
import unicodedata
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Optional

try:
    import anthropic
except ImportError:
    print("❌ anthropic non installé. conda run -n interview python chapter.py ...", file=sys.stderr)
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Cm, Mm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ python-docx non installé.", file=sys.stderr)
    sys.exit(1)


# ── Configuration ─────────────────────────────────────────────────────

PROJECT_DIR   = Path(__file__).parent
CATALOG_PATH  = PROJECT_DIR / "catalog.json"
TRANSCRIPTS_DIR = PROJECT_DIR / "transcripts"
ARTICLES_DIR  = PROJECT_DIR / "articles"
CHAPTERS_DIR  = PROJECT_DIR / "chapters"
LOGO_PATH     = PROJECT_DIR / "assets" / "logo-left.png"

OPUS_MODEL    = "claude-opus-4-5"         # substantive reasoning
SONNET_MODEL  = "claude-sonnet-4-20250514"  # bulk relevance screening

# Collections dont l'auteur est Bhante Bodhidhamma
BHANTE_COLLECTIONS = [
    "guided_meditations",
    "dhammabytes",
    "foundation_course",
    "dharma_talks",
    "international_talks",
    "youtube_channel",
    "bhante_essays",
    "tips_of_the_day",
    "retreat_talks",
]

# Défauts de pipeline
RELEVANCE_THRESHOLD = 6       # score minimum pour entrer dans le pool
TOP_CANDIDATES      = 10      # nombre de candidats full-text envoyés à Opus pour sélection de base
RELEVANCE_BATCH     = 25      # épisodes par appel Sonnet en pass 1
MAX_CONTENT_CHARS   = 100000  # tronque les sources trop longues
MIN_LITE_SCORE      = 45      # pour entrer dans la shortlist audio

# Couleur RGB pour le DOCX
COLOR_BASE = RGBColor(0x00, 0x00, 0x00)


# ── Dataclasses ───────────────────────────────────────────────────────

@dataclass
class Episode:
    collection: str              # catalog key, ex: "dharma_talks"
    slug: str                    # filesystem slug, ex: "dharma-talks"
    season: int
    episode: int
    stem: str
    title: str
    content_type: str            # "audio" ou "text"
    description_short: str
    description_long: str
    lite_score: Optional[int] = None
    duration_seconds: Optional[float] = None
    word_count: Optional[int] = None
    transcript_path: Optional[str] = None
    article_path: Optional[str] = None

    @property
    def source_id(self) -> str:
        return f"{self.collection}/{self.stem}"

    @property
    def short_label(self) -> str:
        return f"{self.title} ({self.collection}, S{self.season:02d}E{self.episode:02d})"

    def content_path(self) -> Optional[Path]:
        if self.content_type == "text":
            if self.article_path:
                return PROJECT_DIR / self.article_path
            return None
        # audio : .txt à côté du .srt
        if self.transcript_path:
            txt = self.transcript_path.replace(".srt", ".txt")
            return PROJECT_DIR / txt
        return None


# ── Utilitaires ───────────────────────────────────────────────────────

def slugify(s: str) -> str:
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode()
    s = re.sub(r"[^\w\s-]", "", s).strip().lower()
    s = re.sub(r"[-\s]+", "-", s)
    return s or "chapter"


def load_bhante_episodes() -> list[Episode]:
    """Charge tous les épisodes de Bhante depuis catalog.json."""
    with open(CATALOG_PATH) as f:
        catalog = json.load(f)

    episodes = []
    for key in BHANTE_COLLECTIONS:
        col = catalog.get(key)
        if not col:
            continue
        slug = col.get("slug", key.replace("_", "-"))
        content_type = col.get("content_type", "audio")
        for season in col.get("seasons", []):
            for ep in season.get("episodes", []):
                stem = ep.get("stem")
                if not stem and ep.get("transcript_path"):
                    stem = Path(ep["transcript_path"]).stem
                if not stem:
                    continue
                episodes.append(Episode(
                    collection=key,
                    slug=slug,
                    season=ep.get("season_number", season.get("number", 1)),
                    episode=ep.get("episode_number", 0),
                    stem=stem,
                    title=ep.get("title", stem),
                    content_type=content_type,
                    description_short=ep.get("description_short", ""),
                    description_long=ep.get("description_long", ""),
                    lite_score=ep.get("lite_score"),
                    duration_seconds=ep.get("duration_seconds"),
                    word_count=ep.get("word_count"),
                    transcript_path=ep.get("transcript_path"),
                    article_path=(f"articles/{slug}/{stem}.txt" if content_type == "text" else None),
                ))
    return episodes


def lookup_source_url(source_id: str) -> Optional[str]:
    """Retrouve l'URL de la source originale depuis le catalogue."""
    parts = source_id.split("/", 1)
    if len(parts) != 2:
        return None
    col_key, stem = parts
    with open(CATALOG_PATH) as f:
        catalog = json.load(f)
    col = catalog.get(col_key)
    if not col:
        return None
    for season in col.get("seasons", []):
        for ep in season.get("episodes", []):
            ep_stem = ep.get("stem") or ""
            if not ep_stem and ep.get("transcript_path"):
                ep_stem = Path(ep["transcript_path"]).stem
            if ep_stem == stem:
                return ep.get("url") or ep.get("page_url")
    return None


def load_content(ep: Episode, max_chars: int = MAX_CONTENT_CHARS) -> Optional[str]:
    """Charge le contenu textuel d'un épisode (transcript ou article)."""
    p = ep.content_path()
    if not p or not p.exists():
        return None
    text = p.read_text(encoding="utf-8", errors="ignore")
    if len(text) > max_chars:
        print(f"    ⚠ TRONCATURE : {ep.source_id} fait {len(text):,} chars, "
              f"limité à {max_chars:,}", file=sys.stderr)
        text = text[:max_chars] + f"\n\n[…texte tronqué à {max_chars} caractères…]"
    return text


def claude_call(client, model: str, system: str, user: str,
                max_tokens: int = 4096, temperature: float = 0.3):
    """Appel Claude avec retry. Utilise le streaming pour les requêtes longues."""
    use_stream = max_tokens > 16000
    for attempt in range(4):
        try:
            if use_stream:
                chunks = []
                with client.messages.stream(
                    model=model,
                    max_tokens=max_tokens,
                    temperature=temperature,
                    system=system,
                    messages=[{"role": "user", "content": user}],
                ) as stream:
                    for text in stream.text_stream:
                        chunks.append(text)
                return "".join(chunks)
            else:
                resp = client.messages.create(
                    model=model,
                    max_tokens=max_tokens,
                    temperature=temperature,
                    system=system,
                    messages=[{"role": "user", "content": user}],
                )
                return resp.content[0].text
        except anthropic.RateLimitError:
            wait = 30 * (attempt + 1)
            print(f"    ⏳ Rate limit, attente {wait}s...")
            time.sleep(wait)
        except anthropic.APIError as e:
            if attempt < 3:
                print(f"    ⚠ API error ({e}), retry dans 15s...")
                time.sleep(15)
            else:
                raise
    raise RuntimeError("Échec Claude après 4 tentatives")


def extract_json(text: str):
    """Extrait le premier bloc JSON d'une réponse."""
    m = re.search(r"```json\s*\n(.*?)\n```", text, re.DOTALL)
    if m:
        return json.loads(m.group(1))
    m = re.search(r"```\s*\n(.*?)\n```", text, re.DOTALL)
    if m:
        return json.loads(m.group(1))
    # fallback : chercher le premier [ ou {
    for i, c in enumerate(text):
        if c in "[{":
            # on essaye par fin décroissante
            for j in range(len(text), i, -1):
                try:
                    return json.loads(text[i:j])
                except json.JSONDecodeError:
                    continue
    raise ValueError(f"Pas de JSON trouvé dans la réponse :\n{text[:800]}")


# ══════════════════════════════════════════════════════════════════════
# PASS 1 — RELEVANCE
# ══════════════════════════════════════════════════════════════════════

def pass1_relevance(client, episodes: list[Episode], theme: str,
                    out_path: Path) -> dict:
    """Note la pertinence de chaque épisode (0-10) vs le thème."""
    print("=" * 68)
    print("PASS 1 — Relevance screening (Sonnet)")
    print("=" * 68)

    if out_path.exists():
        with open(out_path) as f:
            cached = json.load(f)
        print(f"  ✓ cache chargé : {len(cached)} scores")
    else:
        cached = {}

    pending = [ep for ep in episodes if ep.source_id not in cached]
    print(f"  {len(pending)} épisodes à noter ({len(episodes)} total)")

    system = (
        "You are a Theravāda Buddhist scholar screening a catalog of teachings "
        "by Bhante Bodhidhamma to assemble material for a book chapter. "
        "For each episode given, rate its relevance to the chapter theme on "
        "a 0-10 integer scale, based ONLY on title and descriptions. "
        "0 = unrelated; 3 = mentions the topic in passing; "
        "6 = substantially discusses the topic; "
        "9 = the episode is primarily about this topic. "
        "Return STRICT JSON only."
    )

    n_batches = (len(pending) + RELEVANCE_BATCH - 1) // RELEVANCE_BATCH
    for b in range(n_batches):
        batch = pending[b * RELEVANCE_BATCH:(b + 1) * RELEVANCE_BATCH]
        lines = []
        for i, ep in enumerate(batch):
            short = (ep.description_short or "").replace("\n", " ").strip()
            long_ = (ep.description_long or "").replace("\n", " ").strip()
            if len(long_) > 600:
                long_ = long_[:600] + "…"
            lines.append(
                f"[{i+1}] id={ep.source_id}\n"
                f"    title: {ep.title}\n"
                f"    collection: {ep.collection}\n"
                f"    content_type: {ep.content_type}\n"
                f"    short: {short}\n"
                f"    long: {long_}"
            )
        user = (
            f"Chapter theme: \"{theme}\"\n\n"
            f"Episodes:\n" + "\n\n".join(lines) + "\n\n"
            "Return JSON: "
            "{\"scores\": [{\"id\": \"<source_id>\", \"score\": <int>, \"note\": \"<max 12 words>\"}]}\n"
            "One entry per episode, in order."
        )
        print(f"  [{b+1}/{n_batches}] Scoring {len(batch)} épisodes…", end=" ", flush=True)
        txt = claude_call(client, SONNET_MODEL, system, user, max_tokens=4096, temperature=0.2)
        try:
            data = extract_json(txt)
        except Exception as e:
            print(f"\n    ⚠ JSON parse error : {e}")
            continue
        for entry in data.get("scores", []):
            sid = entry.get("id")
            if sid:
                cached[sid] = {
                    "score": int(entry.get("score", 0)),
                    "note": entry.get("note", ""),
                }
        with open(out_path, "w") as f:
            json.dump(cached, f, indent=2, ensure_ascii=False)
        print(f"✓ ({len(cached)} total)")

    return cached


# ══════════════════════════════════════════════════════════════════════
# PASS 2 — BASE SELECTION
# ══════════════════════════════════════════════════════════════════════

def pass2_select_base(client, episodes: list[Episode], relevance: dict,
                      theme: str, top_n: int, out_path: Path) -> dict:
    """Sélectionne LE discours de base via Opus."""
    print("=" * 68)
    print("PASS 2 — Base selection (Opus)")
    print("=" * 68)

    if out_path.exists():
        with open(out_path) as f:
            result = json.load(f)
        print(f"  ✓ base déjà sélectionnée : {result['source_id']}")
        return result

    ep_by_id = {ep.source_id: ep for ep in episodes}

    # Shortlist : score pertinence ≥ seuil, privilégier audio long avec lite_score élevé
    scored = []
    for sid, rel in relevance.items():
        ep = ep_by_id.get(sid)
        if not ep or rel["score"] < RELEVANCE_THRESHOLD:
            continue
        # score de candidature de base
        base_score = rel["score"] * 10
        if ep.content_type == "audio":
            base_score += 15  # on privilégie l'oral pour la voix
            if ep.lite_score and ep.lite_score >= MIN_LITE_SCORE:
                base_score += ep.lite_score * 0.8
            if ep.duration_seconds:
                # privilégie 15-60 min
                mn = ep.duration_seconds / 60
                if 15 <= mn <= 60:
                    base_score += 10
                elif mn > 60:
                    base_score += 5
        else:
            # textes : peuvent être très bons aussi, mais la voix parlée est préférée comme base
            base_score += ep.lite_score * 0.4 if ep.lite_score else 0
        scored.append((base_score, ep, rel))

    if not scored:
        raise RuntimeError("Aucun candidat ne dépasse le seuil de pertinence")

    scored.sort(key=lambda x: -x[0])
    shortlist = scored[:top_n]
    print(f"  shortlist : {len(shortlist)} candidats")
    for i, (bs, ep, rel) in enumerate(shortlist, 1):
        print(f"    [{i}] {ep.short_label} — rel={rel['score']} lite={ep.lite_score} score={bs:.0f}")

    # Charger leurs textes complets
    print("  chargement des transcripts…")
    blocks = []
    valid = []
    for i, (bs, ep, rel) in enumerate(shortlist, 1):
        content = load_content(ep, max_chars=35000)
        if not content:
            print(f"    ⚠ pas de contenu pour {ep.source_id}, skip")
            continue
        valid.append((ep, rel))
        wc = ep.word_count or len(content.split())
        dur = f"{int(ep.duration_seconds/60)}min" if ep.duration_seconds else ""
        blocks.append(
            f"[{len(valid)}] id={ep.source_id}\n"
            f"    title: {ep.title}\n"
            f"    collection: {ep.collection}\n"
            f"    type: {ep.content_type} {dur}\n"
            f"    lite_score: {ep.lite_score}\n"
            f"    relevance: {rel['score']}/10\n"
            f"    word_count: {wc}\n"
            f"    --- FULL TEXT ---\n{content}\n--- END ---"
        )

    system = (
        "You are a master editor of Buddhist literature, preparing a book "
        "chapter drawn from the teachings of Bhante Bodhidhamma. You are "
        "choosing the single best source to use as the BASE text for the "
        "chapter — the spine upon which complementary material from other "
        "talks will later be grafted."
    )
    user = (
        f"Chapter theme: \"{theme}\"\n\n"
        f"Candidates (one is to become the base):\n\n"
        + "\n\n".join(blocks) +
        "\n\nSelect THE ONE candidate that best combines:\n"
        "  • most complete and coherent coverage of the theme,\n"
        "  • highest literary potential for publication (voice, examples,\n"
        "    anecdotes, depth, flow),\n"
        "  • strong enough structure to serve as a spine for the chapter.\n\n"
        "Return JSON ONLY:\n"
        "{\"chosen_index\": <1-based index>, \"source_id\": \"<id>\", "
        "\"justification\": \"<2-4 sentences>\", "
        "\"summary\": \"<2-sentence summary of the chosen base>\"}"
    )
    print(f"  appel Opus (choix de la base parmi {len(valid)} candidats)…")
    txt = claude_call(client, OPUS_MODEL, system, user, max_tokens=2048, temperature=0.3)
    data = extract_json(txt)
    idx = data["chosen_index"] - 1
    chosen_ep, chosen_rel = valid[idx]
    result = {
        "source_id": chosen_ep.source_id,
        "title": chosen_ep.title,
        "collection": chosen_ep.collection,
        "content_type": chosen_ep.content_type,
        "justification": data.get("justification", ""),
        "summary": data.get("summary", ""),
        "relevance": chosen_rel,
    }
    with open(out_path, "w") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
    print(f"  ✓ base choisie : {chosen_ep.short_label}")
    print(f"    {result['justification']}")
    return result


# ══════════════════════════════════════════════════════════════════════
# PASS 4 — COMPOSITION
# ══════════════════════════════════════════════════════════════════════

def pass4_compose(client, episodes: list[Episode], base: dict,
                  theme: str, out_path: Path) -> dict:
    """Structure le discours de base en chapitre (segments, titre, épigraphe)."""
    print("=" * 68)
    print("PASS 4 — Composing chapter (Opus)")
    print("=" * 68)

    if out_path.exists():
        with open(out_path) as f:
            comp = json.load(f)
        print(f"  ✓ composition déjà présente : {len(comp.get('segments', []))} segments")
        return comp

    ep_by_id = {ep.source_id: ep for ep in episodes}
    base_ep = ep_by_id[base["source_id"]]
    base_content = load_content(base_ep, max_chars=MAX_CONTENT_CHARS)
    if not base_content:
        raise RuntimeError("Base content introuvable")

    system = (
        "You are a master literary editor structuring a polished book chapter "
        "drawn from a single teaching by Bhante Bodhidhamma. Your job is to "
        "take the base talk and organise it as a chapter: choose a title, "
        "optional subtitle, optional epigraph, then output the talk as an "
        "ordered sequence of segments. You never invent content. You never "
        "import material from other sources."
    )
    user = (
        f"Chapter theme: \"{theme}\"\n\n"
        f"=== BASE TEXT ===\n"
        f"Title: {base_ep.title}\n"
        f"Collection: {base_ep.collection}\n\n"
        f"{base_content}\n"
        f"=== END BASE TEXT ===\n\n"
        "TASK: structure this talk as a book chapter in JSON.\n\n"
        "=== RULES ===\n"
        "1. ALL segments come from the base talk. Every segment has "
        "source: \"base\".\n\n"
        "2. Follow the talk's original order of ideas. DO NOT reorder, "
        "compress, or paraphrase.\n\n"
        "3. LIGHT CLEANUP ONLY. You may remove transcription hiccups "
        "(mid-sentence false starts, duplicated words, clear typos). "
        "No paraphrase, no added content.\n\n"
        "4. Use \\n\\n inside a segment for paragraph breaks.\n\n"
        "OUTPUT: STRICT JSON of the form\n"
        "{\n"
        "  \"title\": \"<chapter title>\",\n"
        "  \"subtitle\": \"<optional subtitle or empty>\",\n"
        "  \"epigraph\": \"<optional short quote or empty>\",\n"
        "  \"segments\": [\n"
        "    {\"source\": \"base\", \"text\": \"<text, may include \\n\\n>\"},\n"
        "    …\n"
        "  ]\n"
        "}"
    )

    # max_tokens proportionnel au contenu (le JSON de sortie reproduit ~tout le texte)
    content_tokens_est = len(base_content) // 3  # ~3 chars/token
    compose_max_tokens = max(16000, content_tokens_est + 4000)
    print(f"  appel Opus — composition finale… (max_tokens={compose_max_tokens})")
    txt = claude_call(client, OPUS_MODEL, system, user,
                      max_tokens=compose_max_tokens, temperature=0.4)
    comp = extract_json(txt)
    comp["theme"] = theme
    comp["base_source_id"] = base["source_id"]
    # stats
    total_chars = sum(len(s.get("text", "")) for s in comp.get("segments", []))
    comp["stats"] = {"total_chars": total_chars}
    with open(out_path, "w") as f:
        json.dump(comp, f, indent=2, ensure_ascii=False)
    print(f"  ✓ {len(comp.get('segments', []))} segments, {total_chars} chars")
    return comp


# ══════════════════════════════════════════════════════════════════════
# PASS 4c — LIGHT PRUNING (remove superfluous bits)
# ══════════════════════════════════════════════════════════════════════

def pass4c_prune(client, comp: dict, theme: str, out_path: Path) -> dict:
    """Supprime les parties superflues (digressions, boucles orales, redites).

    Opère sur le niveau paragraphe. Conserve la grande majorité du texte."""
    print("=" * 68)
    print("PASS 4c — Light pruning (Opus)")
    print("=" * 68)

    if out_path.exists():
        with open(out_path) as f:
            pruned = json.load(f)
        print(f"  ✓ pruned déjà présent : {len(pruned.get('segments', []))} segments")
        return pruned

    # Aplati la composition en paragraphes
    paragraphs = _extract_flow_paragraphs(comp)
    if not paragraphs:
        raise RuntimeError("Composition vide")

    # Numérote les paragraphes
    numbered = "\n\n".join(f"[{i+1}] {p}" for i, p in enumerate(paragraphs))

    system = (
        "You are a careful book editor performing a single light pruning "
        "pass on a chapter drawn from a Bhante Bodhidhamma dharma talk. "
        "Your ONLY job is to remove genuinely superfluous material "
        "— digressions, redundant restatements, oral loops, filler setups "
        "— while preserving every bit of real substance: doctrine, "
        "anecdotes, examples, metaphors, direct address, voice, "
        "quotations, Pali terms. You apply a LIGHT touch: the chapter "
        "must remain recognisably the same, just tighter. When in doubt, "
        "KEEP."
    )

    user = (
        f"Chapter theme: \"{theme}\"\n\n"
        "=== CURRENT CHAPTER — numbered paragraphs ===\n"
        f"{numbered}\n"
        "=== END CHAPTER ===\n\n"
        "TASK: decide for each numbered paragraph: keep or remove.\n\n"
        "=== REMOVAL CRITERIA (use sparingly — default is KEEP) ===\n"
        "Only remove a paragraph if it clearly matches at least one of:\n"
        "  (a) REDUNDANT RESTATEMENT — a point already made clearly in a "
        "      previous kept paragraph, now being repeated without adding "
        "      nuance;\n"
        "  (b) ORAL LOOP — Bhante circling back to the same thought he "
        "      already expressed, typical of spoken delivery;\n"
        "  (c) DIGRESSION — a tangent that leaves the chapter's "
        "      argumentative thread and adds nothing substantive;\n"
        "  (d) LONG-WINDED SETUP — verbose preamble to a point whose "
        "      actual expression is elsewhere;\n"
        "  (e) PARENTHETICAL ASIDE — a short aside whose removal "
        "      improves flow and loses nothing substantive.\n\n"
        "NEVER REMOVE, UNDER ANY CIRCUMSTANCE:\n"
        "  - anecdotes, personal stories, examples, metaphors, images;\n"
        "  - doctrinal substance introduced for the first time;\n"
        "  - Bhante's voice moments — humour, warmth, direct address;\n"
        "  - concrete illustrations of the theme;\n"
        "  - quotations, Pali or Sanskrit terms, named references;\n"
        "  - any paragraph whose removal would break the argumentative "
        "    flow between neighbours.\n\n"
        "=== VOLUME TARGETS ===\n"
        "Remove about 5-12% of the chapter (by character count). Do NOT "
        "remove more unless the material is truly superfluous. After this "
        "pass, the reader should feel the chapter is tighter but not "
        "shorter in substance.\n\n"
        "=== COHERENCE TEST ===\n"
        "After pruning, mentally read the surviving paragraphs in order. "
        "Every removal must pass the test: the paragraph immediately "
        "before and the paragraph immediately after should still read as "
        "one continuous thought.\n\n"
        "OUTPUT — STRICT JSON:\n"
        "{\n"
        "  \"decisions\": [\n"
        "    {\"idx\": 1, \"action\": \"keep\"},\n"
        "    {\"idx\": 2, \"action\": \"remove\", \"why\": \"<1 sentence>\"},\n"
        "    …one entry per numbered paragraph, in order…\n"
        "  ]\n"
        "}\n"
        "You MUST return exactly one entry per input paragraph."
    )

    print(f"  appel Opus — pruning de {len(paragraphs)} paragraphes…")
    txt = claude_call(client, OPUS_MODEL, system, user,
                      max_tokens=16000, temperature=0.2)
    data = extract_json(txt)
    decisions = data.get("decisions", [])
    dec_by_idx = {d.get("idx", 0): d for d in decisions}

    new_segments = []
    removed = []
    for i, para in enumerate(paragraphs, 1):
        d = dec_by_idx.get(i, {"action": "keep"})
        action = d.get("action", "keep")
        if action == "remove":
            removed.append({
                "idx": i,
                "why": d.get("why", ""),
                "preview": para[:140],
            })
            continue
        # keep
        new_segments.append({"source": "base", "text": para})

    pruned = dict(comp)
    pruned["segments"] = new_segments
    pruned["pruning"] = {
        "original_paragraphs": len(paragraphs),
        "kept_paragraphs": len(new_segments),
        "removed_count": len(removed),
        "removed": removed,
    }
    # stats
    total_chars = sum(len(s.get("text", "")) for s in new_segments)
    original_chars = sum(len(p) for p in paragraphs)
    delta = round(100 * (total_chars - original_chars) / max(original_chars, 1), 1)
    pruned["stats"] = {
        "total_chars": total_chars,
        "original_chars": original_chars,
        "delta_pct": delta,
    }

    with open(out_path, "w") as f:
        json.dump(pruned, f, indent=2, ensure_ascii=False)

    print(f"  ✓ kept {pruned['pruning']['kept_paragraphs']}  "
          f"removed {len(removed)}")
    print(f"  delta: {delta:+}% ({original_chars} → {total_chars} chars)")
    if removed:
        print("  — removals (first 5):")
        for r in removed[:5]:
            print(f"    ✗ #{r['idx']}: {r['why']}")
            print(f"        « {r['preview'][:100]}… »")
    return pruned


# ══════════════════════════════════════════════════════════════════════
# PASS 5 — DOCX RENDERING
# ══════════════════════════════════════════════════════════════════════

def pass5_render_docx(comp: dict, base: dict, theme: str,
                      out_path: Path, source_url: Optional[str] = None) -> None:
    """Rend le chapitre en DOCX (texte noir uniforme)."""
    print("=" * 68)
    print("PASS 5 — Rendering DOCX")
    print("=" * 68)

    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Style par défaut
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(12)

    # Titre
    title = comp.get("title", "Untitled Chapter")
    h = doc.add_heading(title, level=0)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    if comp.get("subtitle"):
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(comp["subtitle"])
        run.italic = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Métadonnées
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Thème : {theme}")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Source originale
    src_p = doc.add_paragraph()
    src_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = src_p.add_run(f"Source : {base.get('title', '')}")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    if source_url:
        src_p.add_run("\n")
        r2 = src_p.add_run(source_url)
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x33, 0x66, 0x99)

    doc.add_paragraph()

    # Épigraphe
    if comp.get("epigraph"):
        ep_p = doc.add_paragraph()
        ep_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = ep_p.add_run(f"« {comp['epigraph']} »")
        r.italic = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_paragraph()

    # Corps : un paragraphe par segment (chaque segment peut contenir \n\n)
    for seg in comp.get("segments", []):
        text = seg.get("text", "")
        if not text.strip():
            continue
        paragraphs = [p for p in text.split("\n\n") if p.strip()]
        for para_text in paragraphs:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.7)
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(para_text.strip())
            run.font.color.rgb = COLOR_BASE

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"  ✓ DOCX → {out_path}")


# ══════════════════════════════════════════════════════════════════════
# PASS 5b — PREPRINT DOCX (typo soignée, logo, prêt pour l'A4)
# ══════════════════════════════════════════════════════════════════════

def _set_cell_borders(cell, colour="999999", sz="4"):
    """Pose une bordure fine autour d'une cellule de tableau."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), sz)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), colour)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _extract_flow_paragraphs(comp: dict) -> list[str]:
    """Extrait le flux texte continu d'une composition en paragraphes."""
    segments = comp.get("segments", [])
    flow_texts: list[str] = []
    for seg in segments:
        t = (seg.get("text") or "").strip()
        if t:
            flow_texts.append(t)
    full_text = "\n\n".join(flow_texts)
    return [p.strip() for p in full_text.split("\n\n") if p.strip()]


def _build_preprint_doc(title: str, subtitle: str, epigraph: str, theme: str,
                        paragraphs: list[str], out_path: Path,
                        edition_label: Optional[str] = None,
                        source_title: Optional[str] = None,
                        source_url: Optional[str] = None) -> None:
    """Construit un DOCX preprint (typo soignée, logo, A4) à partir de paragraphes."""
    INK        = RGBColor(0x1A, 0x1A, 0x1A)
    SOFT_INK   = RGBColor(0x44, 0x44, 0x44)
    MUTED      = RGBColor(0x88, 0x88, 0x88)
    ACCENT     = RGBColor(0x6B, 0x4A, 0x10)  # or sombre

    doc = Document()

    # Format A4, marges généreuses
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(3.2)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.2)
        section.right_margin = Cm(3.2)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

    # Typo de corps — serif élégante
    normal = doc.styles["Normal"]
    normal.font.name = "Garamond"
    normal.font.size = Pt(12)
    # Fallback pour les systèmes sans Garamond (Word pick le plus proche)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), "Garamond")
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(0)

    # ── Page de titre ─────────────────────────────────────────
    # Espace en haut
    for _ in range(3):
        doc.add_paragraph()

    # Logo centré
    if LOGO_PATH.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo_p.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(3.2))

    # Nom du lieu
    place_p = doc.add_paragraph()
    place_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = place_p.add_run("SATIPANYA BUDDHIST TRUST")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    r.bold = True
    # letter-spacing approximatif via small caps désactivé — garde juste majuscules
    place_p.paragraph_format.space_after = Pt(6)

    # Filet supérieur
    rule_top = doc.add_paragraph()
    rule_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rule_top.add_run("·   ·   ·")
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED
    rule_top.paragraph_format.space_before = Pt(6)
    rule_top.paragraph_format.space_after = Pt(18)

    # Chapeau
    hat_p = doc.add_paragraph()
    hat_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hat_p.add_run(f"Teachings of Bhante Bodhidhamma")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = SOFT_INK
    hat_p.paragraph_format.space_after = Pt(36)

    # Titre principal
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(title)
    r.font.size = Pt(28)
    r.font.color.rgb = INK
    r.bold = False
    title_p.paragraph_format.space_after = Pt(6)

    # Sous-titre
    if subtitle:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub_p.add_run(subtitle)
        r.italic = True
        r.font.size = Pt(14)
        r.font.color.rgb = SOFT_INK
        sub_p.paragraph_format.space_after = Pt(12)

    # Filet décoratif
    ornament = doc.add_paragraph()
    ornament.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ornament.add_run("❖")
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT
    ornament.paragraph_format.space_before = Pt(18)
    ornament.paragraph_format.space_after = Pt(18)

    # Épigraphe
    if epigraph:
        ep_p = doc.add_paragraph()
        ep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ep_p.paragraph_format.left_indent = Cm(2.5)
        ep_p.paragraph_format.right_indent = Cm(2.5)
        r = ep_p.add_run(f"« {epigraph} »")
        r.italic = True
        r.font.size = Pt(11)
        r.font.color.rgb = SOFT_INK
        ep_p.paragraph_format.space_after = Pt(6)

    # Theme (petit, discret)
    theme_p = doc.add_paragraph()
    theme_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = theme_p.add_run(f"on the theme of  {theme.lower()}")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    r.italic = True

    # Label d'édition optionnel (ex: "Refined edition")
    if edition_label:
        edl = doc.add_paragraph()
        edl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = edl.add_run(edition_label)
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED
        r.italic = True
        edl.paragraph_format.space_before = Pt(6)

    # Saut de page vers le corps
    doc.add_page_break()

    # ── Entête de chapitre (sur la première page du corps) ────
    # Filet léger au-dessus
    header_rule = doc.add_paragraph()
    header_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header_rule.add_run("―  ―  ―")
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED
    header_rule.paragraph_format.space_before = Pt(12)
    header_rule.paragraph_format.space_after = Pt(12)

    # Titre reprise
    h_p = doc.add_paragraph()
    h_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h_p.add_run(title)
    r.font.size = Pt(20)
    r.font.color.rgb = INK
    h_p.paragraph_format.space_after = Pt(6)

    if subtitle:
        sh_p = doc.add_paragraph()
        sh_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sh_p.add_run(subtitle)
        r.italic = True
        r.font.size = Pt(12)
        r.font.color.rgb = SOFT_INK
        sh_p.paragraph_format.space_after = Pt(6)

    # Petit ornement
    orn2 = doc.add_paragraph()
    orn2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = orn2.add_run("❖")
    r.font.size = Pt(12)
    r.font.color.rgb = ACCENT
    orn2.paragraph_format.space_before = Pt(6)
    orn2.paragraph_format.space_after = Pt(18)

    # ── Corps du chapitre ─────────────────────────────────────
    # Un paragraphe par entrée de la liste, texte noir uniforme.
    for i, para_text in enumerate(paragraphs):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(0)
        if i == 0:
            # premier paragraphe sans indentation, petite capitale
            p.paragraph_format.first_line_indent = Cm(0)
            # Lettrine simplifiée : premier mot plus grand
            first = para_text[0]
            rest = para_text[1:]
            r = p.add_run(first)
            r.font.size = Pt(22)
            r.font.color.rgb = ACCENT
            r.bold = False
            r2 = p.add_run(rest)
            r2.font.size = Pt(12)
            r2.font.color.rgb = INK
        else:
            p.paragraph_format.first_line_indent = Cm(0.8)
            r = p.add_run(para_text)
            r.font.size = Pt(12)
            r.font.color.rgb = INK

    # Ornement final
    doc.add_paragraph()
    end_orn = doc.add_paragraph()
    end_orn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end_orn.add_run("❖   ❖   ❖")
    r.font.size = Pt(12)
    r.font.color.rgb = ACCENT
    end_orn.paragraph_format.space_before = Pt(18)
    end_orn.paragraph_format.space_after = Pt(12)

    # Colophon
    coloph = doc.add_paragraph()
    coloph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = coloph.add_run(
        "Drawn from the teachings of Bhante Bodhidhamma\n"
        "Satipanya Buddhist Trust"
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    # Source originale
    if source_title:
        doc.add_paragraph()
        src_p = doc.add_paragraph()
        src_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = src_p.add_run(f"Based on: {source_title}")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED
        if source_url:
            src_p.add_run("\n")
            r2 = src_p.add_run(source_url)
            r2.font.size = Pt(8)
            r2.font.color.rgb = RGBColor(0x55, 0x77, 0x99)

    # Pied de page : numéro de page centré
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # champ numéro de page
        run = fp.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(fld_end)
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def pass5b_render_preprint(comp: dict, base: dict, theme: str,
                           out_path: Path,
                           source_url: Optional[str] = None) -> None:
    """Rend une version prête à publier — typographie soignée, logo."""
    print("=" * 68)
    print("PASS 5b — Rendering preprint DOCX (publication-ready)")
    print("=" * 68)
    paragraphs = _extract_flow_paragraphs(comp)
    _build_preprint_doc(
        title=comp.get("title", "Untitled Chapter"),
        subtitle=comp.get("subtitle") or "",
        epigraph=comp.get("epigraph") or "",
        theme=theme,
        paragraphs=paragraphs,
        out_path=out_path,
        source_title=base.get("title", ""),
        source_url=source_url,
    )
    print(f"  ✓ preprint DOCX → {out_path}")


# ══════════════════════════════════════════════════════════════════════
# PASS 6 — LIGHT LITERARY REFINEMENT (oral → written, preserves voice)
# ══════════════════════════════════════════════════════════════════════

# Chunking du texte pour rester sous max_tokens et préserver la localité
REFINE_CHUNK_PARAS = 18        # ~paragraphes par appel
REFINE_CONTEXT_PARAS = 3       # paragraphes précédents en contexte (non réécrits)

# Pass 7 — heavier-touch literary refinement
HEAVIER_CHUNK_PARAS = 12       # lots plus petits (travail éditorial plus profond)
HEAVIER_CONTEXT_PARAS = 4      # plus de contexte pour la fluidité

REFINE_SYSTEM = (
    "You are a master literary editor preparing a Buddhist dharma talk for "
    "publication in a printed book. Your SINGLE task is a very light "
    "literary polish: transform oral-register prose into written-register "
    "prose while preserving as much of the original wording as possible. "
    "You apply surgical edits only. You NEVER paraphrase, NEVER add, NEVER "
    "compress, NEVER reorganise. You EXACTLY preserve all anecdotes, "
    "examples, doctrinal content, quotations, Pali terms, and the teacher's "
    "warm, direct, conversational voice. The reader should still hear Bhante "
    "Bodhidhamma — simply as if he were writing, not speaking."
)

HEAVIER_SYSTEM = (
    "You are a senior literary editor preparing the final manuscript of a book "
    "of Buddhist teachings by Bhante Bodhidhamma. This is a DEEPER editorial "
    "pass — the text has already been lightly copy-edited. Your task now is to "
    "bring it to full publication quality: prose that reads as polished written "
    "English while preserving the teacher's warmth, directness, and personality.\n\n"
    "The reader should feel they are reading a book — not a transcript — yet "
    "still recognise Bhante's distinctive voice: his humour, his directness, "
    "his way of making profound teachings feel personal and accessible."
)


def _refine_chunk(client, theme: str, context_paras: list[str],
                  target_paras: list[str]) -> list[str]:
    """Appelle Opus pour polir un lot de paragraphes."""
    # Numérotation stricte pour garantir l'alignement 1-to-1
    target_block = "\n\n".join(
        f"[{i+1}] {p}" for i, p in enumerate(target_paras)
    )
    context_block = ""
    if context_paras:
        context_block = (
            "=== CONTEXT (previous paragraphs, already refined — "
            "DO NOT rewrite, only use for flow and pronoun resolution) ===\n"
            + "\n\n".join(context_paras)
            + "\n=== END CONTEXT ===\n\n"
        )

    user = (
        f"Chapter theme: \"{theme}\"\n\n"
        + context_block +
        "=== PARAGRAPHS TO REFINE ===\n"
        f"{target_block}\n"
        "=== END PARAGRAPHS TO REFINE ===\n\n"
        "TASK: produce a lightly polished version of EACH numbered paragraph.\n\n"
        "ALLOWED edits (use only where they clearly improve readability):\n"
        "  • remove oral fillers: \"you know\", \"I mean\", \"sort of\", "
        "    \"kind of\", \"so\", \"well\", \"right?\", \"OK\" when pure filler;\n"
        "  • remove false starts, self-corrections, mid-sentence restarts, "
        "    dangling repetitions;\n"
        "  • lightly mend grammar and punctuation; close run-on sentences "
        "    with appropriate punctuation; split oral marathon-sentences "
        "    where a clean break exists;\n"
        "  • resolve ambiguous pronouns by restating the antecedent if the "
        "    original is confusing;\n"
        "  • replace overly colloquial connectors with slightly more written "
        "    equivalents (\"and so\" → \"thus\"; \"but then\" → \"yet\") "
        "    SPARINGLY, only where clearly helpful;\n"
        "  • capitalise sentence starts, fix obvious transcription slips;\n"
        "  • RESTORE AUTOMATIC-SPEECH-RECOGNITION ERRORS. This transcript "
        "    was produced by WhisperX from an audio recording and contains "
        "    homophone and word-boundary errors. You are EXPLICITLY "
        "    authorised — and expected — to restore the word the speaker "
        "    clearly intended whenever a phrase is nonsensical or "
        "    grammatically wrong because of such an error. Common patterns:\n"
        "       – homophones: \"side\" ↔ \"aside\", \"affect\" ↔ \"effect\", "
        "         \"there\" ↔ \"their\" ↔ \"they're\", \"to\" ↔ \"too\" "
        "         ↔ \"two\", \"its\" ↔ \"it's\", \"peace\" ↔ \"piece\", "
        "         \"principle\" ↔ \"principal\", \"wander\" ↔ \"wonder\";\n"
        "       – dropped leading letters: \"side\" → \"aside\", \"part\" → "
        "         \"apart\", \"mongst\" → \"amongst\";\n"
        "       – split / merged words: \"a part\" ↔ \"apart\", \"in to\" ↔ "
        "         \"into\", \"any more\" ↔ \"anymore\";\n"
        "       – near-misses for Pali/Sanskrit and proper names that make "
        "         no sense as-written (but be conservative — only if you "
        "         are confident);\n"
        "       – missing small words (articles, prepositions) where the "
        "         sentence is ungrammatical without them.\n"
        "    Apply this restoration SILENTLY as part of the polish. Only "
        "    correct when you are CONFIDENT of the intended word from "
        "    context. When unsure, leave untouched.\n\n"
        "FORBIDDEN:\n"
        "  ✗ paraphrasing or rewording beyond minimal cleanup;\n"
        "  ✗ adding new content, examples, or explanation;\n"
        "  ✗ removing anecdotes, examples, metaphors, direct address, humour;\n"
        "  ✗ compressing or summarising;\n"
        "  ✗ changing the order of sentences or ideas;\n"
        "  ✗ merging or splitting paragraphs (output must have EXACTLY the "
        "    same count as the input — one refined paragraph per numbered "
        "    input paragraph);\n"
        "  ✗ making the voice stiff, academic, formal, or impersonal.\n\n"
        "TARGET FEEL: a gentle copy-edit that feels like a single pass of a "
        "careful friend preparing the text for print. Most sentences should "
        "remain IDENTICAL or near-identical to the original. On average, "
        "expect to change 10-20% of the wording, not more.\n\n"
        "OUTPUT — STRICT JSON, no prose, no markdown:\n"
        "{\"paragraphs\": [\n"
        "  {\"n\": 1, \"text\": \"<refined text of paragraph 1>\"},\n"
        "  {\"n\": 2, \"text\": \"...\"},\n"
        "  ...\n"
        "]}\n"
        "You MUST return exactly one entry per numbered input paragraph, in "
        "the same order."
    )

    txt = claude_call(client, OPUS_MODEL, REFINE_SYSTEM, user,
                      max_tokens=16000, temperature=0.2)
    data = extract_json(txt)
    refined_list = data.get("paragraphs", [])
    # Tri par n au cas où
    refined_list.sort(key=lambda x: x.get("n", 0))
    out = []
    for i, p in enumerate(target_paras):
        if i < len(refined_list):
            t = (refined_list[i].get("text") or "").strip()
            out.append(t if t else p)
        else:
            out.append(p)
    return out


def pass6_refine(client, comp: dict, theme: str, out_path: Path) -> dict:
    """Polit légèrement le chapitre vers un registre écrit, paragraphe par
    paragraphe, en préservant autant que possible le texte original."""
    print("=" * 68)
    print("PASS 6 — Light literary refinement (Opus)")
    print("=" * 68)

    if out_path.exists():
        with open(out_path) as f:
            cached = json.load(f)
        print(f"  ✓ refined.json déjà présent ({len(cached.get('paragraphs', []))} paragraphes)")
        return cached

    paragraphs = _extract_flow_paragraphs(comp)
    print(f"  {len(paragraphs)} paragraphes à polir")

    refined: list[str] = []
    i = 0
    batch_num = 0
    while i < len(paragraphs):
        chunk = paragraphs[i:i + REFINE_CHUNK_PARAS]
        # Contexte : les 3 derniers paragraphes déjà polis
        context = refined[-REFINE_CONTEXT_PARAS:] if refined else []
        batch_num += 1
        print(f"  [{batch_num}] polissage {len(chunk)} paragraphes "
              f"({i+1}→{i+len(chunk)}/{len(paragraphs)})…", end=" ", flush=True)
        try:
            polished = _refine_chunk(client, theme, context, chunk)
            refined.extend(polished)
            print("✓")
        except Exception as e:
            print(f"⚠ {e} — on garde le texte original")
            refined.extend(chunk)
        # Save partial
        partial = {
            "theme": theme,
            "title": comp.get("title", ""),
            "subtitle": comp.get("subtitle", ""),
            "epigraph": comp.get("epigraph", ""),
            "paragraphs": refined,
            "source_paragraph_count": len(paragraphs),
            "progress": f"{len(refined)}/{len(paragraphs)}",
        }
        with open(out_path, "w") as f:
            json.dump(partial, f, indent=2, ensure_ascii=False)
        i += REFINE_CHUNK_PARAS

    # Stats de similarité approximative
    original_chars = sum(len(p) for p in paragraphs)
    refined_chars = sum(len(p) for p in refined)
    delta_pct = round(100 * (refined_chars - original_chars) / max(original_chars, 1), 1)
    print(f"  ✓ refined : {refined_chars} chars ({delta_pct:+}% vs original {original_chars})")

    result = {
        "theme": theme,
        "title": comp.get("title", ""),
        "subtitle": comp.get("subtitle", ""),
        "epigraph": comp.get("epigraph", ""),
        "paragraphs": refined,
        "source_paragraph_count": len(paragraphs),
        "stats": {
            "original_chars": original_chars,
            "refined_chars": refined_chars,
            "delta_pct": delta_pct,
        },
    }
    with open(out_path, "w") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
    return result


def pass6b_render_refined(refined: dict, theme: str, out_path: Path,
                          source_title: Optional[str] = None,
                          source_url: Optional[str] = None) -> None:
    """Rend la version refined via le même gabarit preprint."""
    print("=" * 68)
    print("PASS 6b — Rendering refined preprint DOCX")
    print("=" * 68)
    _build_preprint_doc(
        title=refined.get("title", "Untitled Chapter"),
        subtitle=refined.get("subtitle") or "",
        epigraph=refined.get("epigraph") or "",
        theme=theme,
        paragraphs=refined.get("paragraphs", []),
        out_path=out_path,
        edition_label="Refined edition · lightly polished for print",
        source_title=source_title,
        source_url=source_url,
    )
    print(f"  ✓ refined preprint DOCX → {out_path}")


# ── Pass 7 — Heavier-touch literary refinement ───────────────────────

def _heavier_refine_chunk(client, theme: str, context_paras: list[str],
                          target_paras: list[str]) -> list[str]:
    """Appelle Opus pour un polissage éditorial plus profond."""
    target_block = "\n\n".join(
        f"[{i+1}] {p}" for i, p in enumerate(target_paras)
    )
    context_block = ""
    if context_paras:
        context_block = (
            "=== CONTEXT (previous paragraphs, already edited — "
            "DO NOT rewrite, only use for flow and continuity) ===\n"
            + "\n\n".join(context_paras)
            + "\n=== END CONTEXT ===\n\n"
        )

    user = (
        f"Chapter theme: \"{theme}\"\n\n"
        + context_block +
        "=== PARAGRAPHS TO EDIT ===\n"
        f"{target_block}\n"
        "=== END PARAGRAPHS TO EDIT ===\n\n"
        "TASK: produce a fully edited version of EACH numbered paragraph, "
        "bringing it to publication quality.\n\n"
        "ALLOWED edits (use freely where they improve the prose):\n"
        "  • Rewrite sentences for fluency — reshape awkward oral "
        "    constructions into natural written prose;\n"
        "  • Improve paragraph transitions — add a brief linking phrase "
        "    where paragraphs feel disconnected;\n"
        "  • Tighten verbose passages — condense wordy oral expressions "
        "    without losing meaning;\n"
        "  • Vary sentence rhythm — break monotonous patterns, combine "
        "    choppy fragments, split overlong sentences;\n"
        "  • Strengthen paragraph openings and closings;\n"
        "  • Convert oral signposts to written ones (\"As I said earlier\" "
        "    → cut or replace with a more writerly transition);\n"
        "  • All the light-touch edits from the previous pass remain "
        "    available: fix grammar, remove residual fillers, restore "
        "    ASR errors, resolve ambiguous pronouns.\n\n"
        "VOICE PRESERVATION (critical):\n"
        "  • The speaker's own vocabulary is the CEILING — restructure "
        "    sentences freely, but reuse his words. Do NOT upgrade plain "
        "    words to fancier synonyms (\"much bigger\" stays \"much bigger\", "
        "    not \"considerably larger\"; \"born evil\" stays \"born evil\", "
        "    not \"born sinful\");\n"
        "  • KEEP contractions — Bhante speaks with contractions (\"I can't\", "
        "    \"it's\", \"don't\", \"we've\") and expanding them makes the "
        "    voice stiff. Never change \"can't\" to \"cannot\" etc.;\n"
        "  • Do NOT add metaphors, imagery, or poetic flourishes the speaker "
        "    did not use — you may rearrange his images, not invent new ones;\n"
        "  • Preserve his characteristic phrases: \"shall we say\", direct "
        "    questions to the audience, dry understatement, informal asides.\n\n"
        "FORBIDDEN:\n"
        "  ✗ Adding new doctrinal content, examples, or arguments;\n"
        "  ✗ Removing anecdotes, stories, metaphors, humour, Pali terms;\n"
        "  ✗ Changing the order of ideas or merging/splitting paragraphs "
        "    (output must have EXACTLY the same count as the input — one "
        "    edited paragraph per numbered input paragraph);\n"
        "  ✗ Making the voice stiff, academic, or impersonal;\n"
        "  ✗ Upgrading vocabulary — the text should read at the SAME register "
        "    as the original, not a higher one.\n\n"
        "TARGET FEEL: \"a talk so good it reads like an essay\" — the prose "
        "should flow like polished written English while preserving Bhante's "
        "warmth, directness, and personality. The improvement comes from "
        "STRUCTURE (sentence flow, rhythm, tightening) not from DICTION "
        "(fancier words). Expect 25-40% wording change compared to the "
        "input.\n\n"
        "OUTPUT — STRICT JSON, no prose, no markdown:\n"
        "{\"paragraphs\": [\n"
        "  {\"n\": 1, \"text\": \"<edited text of paragraph 1>\"},\n"
        "  {\"n\": 2, \"text\": \"...\"},\n"
        "  ...\n"
        "]}\n"
        "You MUST return exactly one entry per numbered input paragraph, in "
        "the same order."
    )

    txt = claude_call(client, OPUS_MODEL, HEAVIER_SYSTEM, user,
                      max_tokens=16000, temperature=0.4)
    data = extract_json(txt)
    refined_list = data.get("paragraphs", [])
    refined_list.sort(key=lambda x: x.get("n", 0))
    out = []
    for i, p in enumerate(target_paras):
        if i < len(refined_list):
            t = (refined_list[i].get("text") or "").strip()
            out.append(t if t else p)
        else:
            out.append(p)
    return out


def pass7_heavier_refine(client, refined: dict, theme: str, out_path: Path) -> dict:
    """Polissage éditorial plus profond — registre livre publié, en
    préservant la voix de Bhante."""
    print("=" * 68)
    print("PASS 7 — Heavier literary refinement (Opus)")
    print("=" * 68)

    if out_path.exists():
        with open(out_path) as f:
            cached = json.load(f)
        print(f"  ✓ refined_heavier.json déjà présent ({len(cached.get('paragraphs', []))} paragraphes)")
        return cached

    paragraphs = refined.get("paragraphs", [])
    print(f"  {len(paragraphs)} paragraphes à éditer")

    edited: list[str] = []
    i = 0
    batch_num = 0
    while i < len(paragraphs):
        chunk = paragraphs[i:i + HEAVIER_CHUNK_PARAS]
        context = edited[-HEAVIER_CONTEXT_PARAS:] if edited else []
        batch_num += 1
        print(f"  [{batch_num}] édition {len(chunk)} paragraphes "
              f"({i+1}→{i+len(chunk)}/{len(paragraphs)})…", end=" ", flush=True)
        try:
            polished = _heavier_refine_chunk(client, theme, context, chunk)
            edited.extend(polished)
            print("✓")
        except Exception as e:
            print(f"⚠ {e} — on garde le texte du pass 6")
            edited.extend(chunk)
        # Save partial
        partial = {
            "theme": theme,
            "title": refined.get("title", ""),
            "subtitle": refined.get("subtitle", ""),
            "epigraph": refined.get("epigraph", ""),
            "paragraphs": edited,
            "source_paragraph_count": len(paragraphs),
            "progress": f"{len(edited)}/{len(paragraphs)}",
        }
        with open(out_path, "w") as f:
            json.dump(partial, f, indent=2, ensure_ascii=False)
        i += HEAVIER_CHUNK_PARAS

    original_chars = sum(len(p) for p in paragraphs)
    edited_chars = sum(len(p) for p in edited)
    delta_pct = round(100 * (edited_chars - original_chars) / max(original_chars, 1), 1)
    print(f"  ✓ heavier refined : {edited_chars} chars ({delta_pct:+}% vs pass 6 input {original_chars})")

    result = {
        "theme": theme,
        "title": refined.get("title", ""),
        "subtitle": refined.get("subtitle", ""),
        "epigraph": refined.get("epigraph", ""),
        "paragraphs": edited,
        "source_paragraph_count": len(paragraphs),
        "stats": {
            "original_chars": original_chars,
            "edited_chars": edited_chars,
            "delta_pct": delta_pct,
        },
    }
    with open(out_path, "w") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
    return result


def pass7b_render_heavier(refined: dict, theme: str, out_path: Path,
                          source_title: Optional[str] = None,
                          source_url: Optional[str] = None) -> None:
    """Rend la version heavier-refined via le même gabarit preprint."""
    print("=" * 68)
    print("PASS 7b — Rendering heavier-refined preprint DOCX")
    print("=" * 68)
    _build_preprint_doc(
        title=refined.get("title", "Untitled Chapter"),
        subtitle=refined.get("subtitle") or "",
        epigraph=refined.get("epigraph") or "",
        theme=theme,
        paragraphs=refined.get("paragraphs", []),
        out_path=out_path,
        edition_label="Book edition · fully edited for publication",
        source_title=source_title,
        source_url=source_url,
    )
    print(f"  ✓ heavier-refined preprint DOCX → {out_path}")


# ══════════════════════════════════════════════════════════════════════
# ORCHESTRATION
# ══════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Compose a book chapter from Bhante Bodhidhamma's teachings.")
    parser.add_argument("--theme", required=True, help="Chapter theme (ex: 'Dharma in Daily Life')")
    parser.add_argument("--top-candidates", type=int, default=TOP_CANDIDATES,
                        help=f"Candidats full-text pour la sélection de base (def: {TOP_CANDIDATES})")
    parser.add_argument("--output-dir", default=None,
                        help="Dossier de sortie (def: chapters/<slug>)")
    parser.add_argument("--rescreen", action="store_true",
                        help="Force le rescoring de pertinence (pass 1)")
    args = parser.parse_args()

    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("❌ ANTHROPIC_API_KEY non défini", file=sys.stderr)
        sys.exit(1)

    theme = args.theme.strip()
    slug = slugify(theme)
    out_dir = Path(args.output_dir) if args.output_dir else (CHAPTERS_DIR / slug)
    out_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n📖 Chapter composition — theme: « {theme} »")
    print(f"   output dir : {out_dir}\n")

    episodes = load_bhante_episodes()
    print(f"📚 {len(episodes)} épisodes de Bhante chargés")
    audio = sum(1 for e in episodes if e.content_type == "audio")
    text = sum(1 for e in episodes if e.content_type == "text")
    print(f"   ({audio} audio + {text} texte)\n")

    client = anthropic.Anthropic()

    cand_path = out_dir / "candidates.json"
    base_path = out_dir / "base.json"
    comp_path = out_dir / "composition.json"
    pruned_path = out_dir / "pruned.json"
    refined_path = out_dir / "refined.json"
    heavier_path = out_dir / "refined_heavier.json"
    docx_path = out_dir / "chapter.docx"
    preprint_path = out_dir / "chapter-preprint.docx"
    refined_docx_path = out_dir / "chapter-preprint-refined.docx"
    heavier_docx_path = out_dir / "chapter-preprint-refined-heavier.docx"

    if args.rescreen and cand_path.exists():
        cand_path.unlink()

    relevance = pass1_relevance(client, episodes, theme, cand_path)
    base = pass2_select_base(client, episodes, relevance, theme,
                             args.top_candidates, base_path)
    source_url = lookup_source_url(base["source_id"])

    comp = pass4_compose(client, episodes, base, theme, comp_path)
    pruned = pass4c_prune(client, comp, theme, pruned_path)
    pass5_render_docx(pruned, base, theme, docx_path, source_url=source_url)
    pass5b_render_preprint(pruned, base, theme, preprint_path, source_url=source_url)
    refined = pass6_refine(client, pruned, theme, refined_path)
    pass6b_render_refined(refined, theme, refined_docx_path,
                          source_title=base.get("title", ""),
                          source_url=source_url)
    heavier = pass7_heavier_refine(client, refined, theme, heavier_path)
    pass7b_render_heavier(heavier, theme, heavier_docx_path,
                          source_title=base.get("title", ""),
                          source_url=source_url)

    print(f"\n✅ Terminé.")
    print(f"   Chapitre (relecture)            : {docx_path}")
    print(f"   Chapitre (prêt à imprimer)      : {preprint_path}")
    print(f"   Chapitre (prêt à imprimer, poli): {refined_docx_path}")
    print(f"   Chapitre (édition livre)        : {heavier_docx_path}")


if __name__ == "__main__":
    main()
