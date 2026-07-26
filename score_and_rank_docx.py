#!/usr/bin/env python3
"""Score & classe les talks parlés de Satipañña (même double rubrique que les retraites IMS),
puis génère un RAPPORT DOCX — sans rien ajouter au site /site.

Deux rubriques par talk (chacune 0-100 : inspiration, soundness, overall) :
  • full  — soundness = justesse TECHNIQUE + DOCTRINALE (fidélité au Dhamma)
  • craft — IDENTIQUE mais critère doctrinal RETIRÉ : soundness = clarté, rigueur conceptuelle
            et justesse pratique des instructions uniquement.

Les notes sont stockées dans un FICHIER ANNEXE (rankings_scores.json) — le catalog.json du
projet sœur n'est PAS modifié. Le script est résumable.

Rapport DOCX : par collection (2 classements) + combiné (2 classements) + côte-à-côte (Δ).

Usage:
    python score_and_rank_docx.py                 # score le manquant puis génère le DOCX
    python score_and_rank_docx.py --build-only     # régénère seulement le DOCX depuis l'annexe
    python score_and_rank_docx.py --force          # re-scorer tout
    python score_and_rank_docx.py --dry-run        # compter sans appeler l'API
"""

import json
import os
import re
import sys
import time
import argparse
from pathlib import Path

try:
    import anthropic
except ImportError:
    print("ERREUR : pip install anthropic"); sys.exit(1)

PROJECT_DIR = Path(__file__).parent
CATALOG_PATH = PROJECT_DIR / "catalog.json"
ARTICLES_DIR = PROJECT_DIR / "articles"
TRANSCRIPTS_DIR = PROJECT_DIR / "transcripts"
SCORES_PATH = PROJECT_DIR / "rankings_scores.json"      # annexe (pas catalog.json)
DOCX_PATH = PROJECT_DIR / "Satipanya_Talk_Rankings.docx"

CLAUDE_MODEL = "claude-sonnet-4-6"
MAX_TRANSCRIPT_CHARS = 80_000

# Collections de talks PARLÉS dans le scope (exclut méditations guidées + essais texte + tips).
SCOPE_KEYS = [
    "dharma_talks", "youtube_channel", "dhammabytes",
    "international_talks", "noirins_teachings", "retreat_talks", "foundation_course",
]

# ── Rubriques (identiques aux retraites IMS) ───────────────────
_INSPIRATION = """A. INSPIRATION — How compelling and inspirational is it?
   - Does it move and energize the listener toward practice?
   - Memorable images, stories, framings; warmth and authenticity of voice;
   - Does it leave you wanting to sit, to investigate, to live more wisely?"""

_SOUNDNESS_FULL = """B. SOUNDNESS — Is it technically and doctrinally sound?
   - Doctrinal accuracy: faithful to the Buddha's teaching (four noble truths, anattā, dependent
     origination, the brahmavihāras, etc.); no distortion, no sloppy or misleading claims.
   - Technical/practical soundness: are the meditation instructions correct, precise, and safe to follow?
   - Conceptual rigor: distinctions drawn carefully; nuance where nuance is due."""

_SOUNDNESS_CRAFT = """B. CRAFT SOUNDNESS — Is it technically and practically sound, AS CRAFT, independent of doctrine?
   - Technical/practical soundness: are the meditation instructions correct, precise, clearly
     sequenced, and safe to follow?
   - Conceptual rigor & clarity: distinctions drawn carefully; terms used consistently; the talk
     well-structured, coherent, and easy to follow.
   - Practical usefulness: could a practitioner actually act on it?
   IMPORTANT — DO NOT assess fidelity to Buddhist doctrine, scriptural accuracy, sectarian/lineage
   framing, or whether claims align with Theravāda teaching. Ignore doctrinal correctness ENTIRELY:
   a talk that blends traditions, quotes other lineages, or departs from canonical positions must NOT
   be penalized here — only its clarity, craft, and practical soundness as instruction matter."""

_OVERALL = """Then give an OVERALL score 0-100 that REQUIRES BOTH. A talk strong on one dimension but
weak on the other must NOT score high overall. Reward talks that are simultaneously moving AND sound.

CALIBRATION:
- 85-100: genuinely outstanding on both dimensions.
- 65-84 : strong on both, minor weaknesses.
- 45-64 : solid but unremarkable, or strong on one dimension and average on the other.
- 25-44 : weak — flat delivery, thin content, or shaky on the second dimension.
- 0-24  : a short logistical/instructional fragment, mostly guided silence, or seriously unsound.

Respond with ONLY a JSON object, no other text:
{"inspiration": <int 0-100>, "soundness": <int 0-100>, "overall": <int 0-100>, "reason": "<one-sentence justification weighing both dimensions>"}"""

_HEAD = """You are an expert evaluator of Theravāda Buddhist dharma talks, combining two competencies:
a contemplative practitioner-scholar and a discerning listener (what truly inspires).

Evaluate this talk transcript on TWO dimensions, each 0-100:
"""

PROFILES = {
    "full":  {"rubric": f"{_HEAD}\n{_INSPIRATION}\n\n{_SOUNDNESS_FULL}\n\n{_OVERALL}",  "sound": "Sound"},
    "craft": {"rubric": f"{_HEAD}\n{_INSPIRATION}\n\n{_SOUNDNESS_CRAFT}\n\n{_OVERALL}", "sound": "Craft"},
}


# ── Lecture des transcriptions ─────────────────────────────────
def parse_srt(p: Path) -> str:
    out = []
    for line in p.read_text(encoding="utf-8", errors="replace").split("\n"):
        line = line.strip()
        if not line or re.match(r"^\d+$", line) or re.match(r"\d{2}:\d{2}:\d{2}", line):
            continue
        out.append(line)
    return " ".join(out)


def truncate(text: str) -> str:
    if len(text) <= MAX_TRANSCRIPT_CHARS:
        return text
    half = MAX_TRANSCRIPT_CHARS // 2
    return text[:half] + "\n\n[... transcript truncated ...]\n\n" + text[-half:]


def ep_stem(ep: dict) -> str:
    tp = ep.get("transcript_path", "")
    if tp:
        return Path(tp).stem
    return ep.get("stem", "") or ""


def load_text(slug: str, ep: dict) -> str:
    """Préfère l'article embelli, puis le .txt de transcription, puis le .srt."""
    stem = ep_stem(ep)
    tp = ep.get("transcript_path", "")
    candidates = [ARTICLES_DIR / slug / f"{stem}.txt"]
    if tp:
        candidates.append(PROJECT_DIR / tp.replace(".srt", ".txt"))
    for c in candidates:
        if c.exists():
            return c.read_text(encoding="utf-8", errors="replace")
    if tp and (PROJECT_DIR / tp).exists():
        return parse_srt(PROJECT_DIR / tp)
    return ""


# ── Appel Claude ───────────────────────────────────────────────
def score(client, rubric, title, speaker, text) -> dict:
    resp = client.messages.create(
        model=CLAUDE_MODEL, max_tokens=300,
        messages=[{"role": "user",
                   "content": f"{rubric}\n\n# Talk: {title}\n# Teacher: {speaker}\n\n# Transcript:\n\n{text}"}],
    )
    raw = resp.content[0].text.strip()
    m = re.search(r"\{.*\}", raw, re.DOTALL)
    if m:
        try:
            r = json.loads(m.group())
            o = {k: int(r.get(k, -1)) for k in ("inspiration", "soundness", "overall")}
            o["reason"] = str(r.get("reason", ""))
            if all(0 <= o[k] <= 100 for k in ("inspiration", "soundness", "overall")):
                return o
        except (ValueError, json.JSONDecodeError):
            pass
    print(f"  ⚠ réponse inattendue: {raw[:160]}")
    return None


# ── Génération du rapport DOCX ─────────────────────────────────
def build_docx(catalog, scores):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    names = {k: catalog[k].get("name", k) for k in SCOPE_KEYS if k in catalog}

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(0.5))

    base = doc.styles["Normal"].font
    base.name = "Calibri"; base.size = Pt(9)

    def cell_text(cell, text, size=8.5, bold=False, align=None, color=None):
        cell.text = ""
        p = cell.paragraphs[0]
        if align is not None:
            p.alignment = align
        run = p.add_run("" if text is None else str(text))
        run.font.size = Pt(size); run.bold = bold
        if color:
            run.font.color.rgb = color

    GREY = RGBColor(0x70, 0x70, 0x70)
    GREEN = RGBColor(0x2F, 0x85, 0x5A)
    ORANGE = RGBColor(0xC0, 0x56, 0x21)

    def add_table(headers, rows, widths):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.autofit = False
        for c, (htext, w) in enumerate(zip(headers, widths)):
            cell_text(t.rows[0].cells[c], htext, size=8.5, bold=True)
            t.columns[c].width = Inches(w)
        for r in rows:
            cells = t.add_row().cells
            for c, val in enumerate(r):
                cell, color = cells[c], None
                if isinstance(val, tuple):       # (text, color)
                    val, color = val
                align = WD_ALIGN_PARAGRAPH.CENTER if c < len(widths) - 1 and isinstance(val, int) else None
                cell_text(cell, val, size=8, color=color)
                cell.width = Inches(widths[c])
        doc.add_paragraph()

    # Ranked list for one profile (prefix dhamma/craft) within a set of items.
    def ranked_rows(items, prefix, with_collection):
        key = f"{prefix}_overall"
        rk = [it for it in items if it.get(key) is not None]
        rk.sort(key=lambda it: it[key], reverse=True)
        rows = []
        for i, it in enumerate(rk, 1):
            row = [i, it[key], it[f"{prefix}_inspiration"], it[f"{prefix}_soundness"],
                   it["speaker"], it["title"]]
            if with_collection:
                row.append(it["collection"])
            row.append(it[f"{prefix}_reason"])
            rows.append(row)
        return rk, rows

    # ---- Title + methodology ----
    h = doc.add_heading("Satipañña — Talk Rankings", 0)
    doc.add_paragraph(
        "Every spoken dharma talk across seven Satipañña collections (629 talks), scored 0–100 by "
        "Claude (sonnet-4-6) on two dimensions — INSPIRATION (how compelling and inspiring) and "
        "SOUNDNESS — with an OVERALL that requires both. Two rankings are given: one where soundness "
        "includes doctrinal fidelity to the Dhamma, and one where the doctrinal criterion is removed "
        "(“craft”: clarity, conceptual rigor, and practical soundness of the instructions only)."
    )
    doc.add_paragraph(
        "Scores live in a sidecar file (rankings_scores.json); the project catalog and website are untouched."
    ).runs[0].italic = True

    # Flatten all in-scope scored items.
    all_items = [it for it in scores.values() if it.get("full_overall") is not None]

    HDR_COLL = ["#", "Ovr", "Insp", "Snd", "Teacher", "Title", "Why"]
    W_COLL = [0.35, 0.45, 0.5, 0.5, 1.3, 2.2, 5.0]
    HDR_COMB = ["#", "Ovr", "Insp", "Snd", "Teacher", "Title", "Collection"]
    W_COMB = [0.4, 0.5, 0.5, 0.5, 1.6, 4.2, 2.0]

    # ---- Per-collection ----
    doc.add_heading("Part 1 — By collection", 1)
    for k in SCOPE_KEYS:
        items = [it for it in all_items if it["collection_key"] == k]
        if not items:
            continue
        doc.add_heading(f"{names.get(k, k)}  ({len(items)} talks)", 2)

        doc.add_heading("Ranked WITH doctrinal criterion", 3)
        _, rows = ranked_rows(items, "full", with_collection=False)
        add_table(["#", "Ovr", "Insp", "Sound", "Teacher", "Title", "Why"], rows, W_COLL)

        doc.add_heading("Ranked WITHOUT doctrine (craft)", 3)
        _, rows = ranked_rows(items, "craft", with_collection=False)
        add_table(["#", "Ovr", "Insp", "Craft", "Teacher", "Title", "Why"], rows, W_COLL)

    # ---- Combined ----
    doc.add_page_break()
    doc.add_heading("Part 2 — Combined (all 629 talks)", 1)

    doc.add_heading("Combined ranking — WITH doctrine", 2)
    _, rows = ranked_rows(all_items, "full", with_collection=True)
    add_table(HDR_COMB, [r[:7] for r in rows], W_COMB)

    doc.add_heading("Combined ranking — WITHOUT doctrine (craft)", 2)
    _, rows = ranked_rows(all_items, "craft", with_collection=True)
    add_table(["#", "Ovr", "Insp", "Craft", "Teacher", "Title", "Collection"],
              [r[:7] for r in rows], W_COMB)

    # ---- Side by side ----
    doc.add_page_break()
    doc.add_heading("Part 3 — Side by side (craft vs. doctrine)", 1)
    doc.add_paragraph(
        "All talks sorted by craft score. Δ = craft − doctrine: positive (green) means the talk rises "
        "once doctrine is set aside; negative (orange) means doctrinal fidelity was lifting its score."
    )
    sbs = [it for it in all_items if it.get("craft_overall") is not None]
    sbs.sort(key=lambda it: it["craft_overall"], reverse=True)
    rows = []
    for i, it in enumerate(sbs, 1):
        delta = it["craft_overall"] - it["full_overall"]
        col = GREEN if delta > 0 else ORANGE if delta < 0 else GREY
        rows.append([i, it["craft_overall"], it["full_overall"],
                     (f"{'+' if delta > 0 else ''}{delta}", col),
                     it["speaker"], it["title"], it["collection"]])
    add_table(["#", "Craft", "Doctrine", "Δ", "Teacher", "Title", "Collection"],
              rows, [0.4, 0.6, 0.7, 0.5, 1.6, 4.0, 2.0])

    doc.save(DOCX_PATH)
    print(f"\nRapport DOCX écrit : {DOCX_PATH}  ({len(all_items)} talks)")


# ── Main ───────────────────────────────────────────────────────
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--build-only", action="store_true")
    ap.add_argument("--force", action="store_true")
    ap.add_argument("--dry-run", action="store_true")
    args = ap.parse_args()

    with open(CATALOG_PATH, encoding="utf-8") as f:
        catalog = json.load(f)

    scores = {}
    if SCORES_PATH.exists():
        scores = json.loads(SCORES_PATH.read_text(encoding="utf-8"))

    if args.build_only:
        build_docx(catalog, scores)
        return

    # Liste des (id, slug, ep, col_key, col_name).
    todo = []
    for k in SCOPE_KEYS:
        col = catalog.get(k)
        if not col:
            print(f"  ⚠ collection absente : {k}"); continue
        slug = col.get("slug", k.replace("_", "-"))
        cname = col.get("name", k)
        for s in col.get("seasons", []):
            for ep in s.get("episodes", []):
                stem = ep_stem(ep)
                if not stem:
                    continue
                eid = f"{k}/{stem}"
                entry = scores.get(eid, {})
                entry.update({"collection_key": k, "collection": cname,
                              "speaker": ep.get("speaker", ""), "title": ep.get("title", "")})
                scores[eid] = entry
                need = [p for p in PROFILES if args.force or f"{p}_overall" not in entry]
                if need:
                    todo.append((eid, slug, ep, need))

    n_calls = sum(len(n) for _, _, _, n in todo)
    print(f"Talks en scope : {len([1 for k in SCOPE_KEYS for s in catalog.get(k,{}).get('seasons',[]) for _ in s['episodes']])}")
    print(f"Talks à (re)scorer : {len(todo)}  |  appels API : {n_calls}")
    if args.dry_run:
        return
    if not todo:
        build_docx(catalog, scores); return

    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("ERREUR : ANTHROPIC_API_KEY non définie"); sys.exit(1)
    client = anthropic.Anthropic()

    done = errs = 0
    for j, (eid, slug, ep, need) in enumerate(todo):
        entry = scores[eid]
        text = load_text(slug, ep)
        print(f"[{j+1}/{len(todo)}] {entry['collection']} — {entry['speaker']}: {entry['title'][:48]}", flush=True)
        if len(text) < 50:
            for p in need:
                entry[f"{p}_overall"] = entry[f"{p}_inspiration"] = entry[f"{p}_soundness"] = 5
                entry[f"{p}_reason"] = "Transcript missing or too short to evaluate."
            done += 1
        else:
            ttext = truncate(text)
            for p in need:
                try:
                    r = score(client, PROFILES[p]["rubric"], entry["title"], entry["speaker"], ttext)
                    if r:
                        entry[f"{p}_overall"] = r["overall"]
                        entry[f"{p}_inspiration"] = r["inspiration"]
                        entry[f"{p}_soundness"] = r["soundness"]
                        entry[f"{p}_reason"] = r["reason"]
                        print(f"    {p}: {r['overall']}/100 (insp {r['inspiration']} · snd {r['soundness']})")
                        done += 1
                    else:
                        errs += 1
                except Exception as e:
                    print(f"    {p} ✗ {e}"); errs += 1
                time.sleep(0.2)
        SCORES_PATH.write_text(json.dumps(scores, ensure_ascii=False, indent=1), encoding="utf-8")

    print(f"\nScoring terminé : {done} notes, {errs} erreurs")
    build_docx(catalog, scores)


if __name__ == "__main__":
    main()
